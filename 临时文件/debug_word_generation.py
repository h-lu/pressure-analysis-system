#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
调试Word报告生成
"""
import sys
import os
import json
import traceback
from pathlib import Path

# 添加backend到path
sys.path.append(os.path.join(os.path.dirname(__file__), 'backend'))

def debug_word_generation():
    """调试Word报告生成"""
    task_id = "9ea02dd3-9e29-4309-ab3e-1c2172909982"
    
    print(f"🔧 调试Word报告生成，task_id: {task_id}")
    print("=" * 60)
    
    # 1. 检查分析结果文件
    print("1. 检查分析结果文件...")
    charts_dir = Path("backend/static/charts") / task_id
    analysis_file = charts_dir / "analysis_results.json"
    
    if analysis_file.exists():
        print(f"✅ 分析结果文件存在: {analysis_file}")
        print(f"📊 文件大小: {analysis_file.stat().st_size:,} 字节")
        
        # 读取和验证JSON
        try:
            with open(analysis_file, 'r', encoding='utf-8') as f:
                analysis_data = json.load(f)
            print(f"✅ JSON文件格式正确")
            print(f"📋 包含的键: {list(analysis_data.keys())}")
            
            # 检查数据结构
            print("\n📊 数据结构分析:")
            for key, value in analysis_data.items():
                if isinstance(value, list):
                    print(f"  {key}: 列表, 长度={len(value)}")
                    if len(value) > 0:
                        print(f"    第一个元素类型: {type(value[0])}")
                elif isinstance(value, dict):
                    print(f"  {key}: 字典, 键数量={len(value)}")
                else:
                    print(f"  {key}: {type(value)}")
        except Exception as e:
            print(f"❌ JSON文件读取失败: {e}")
            return False
    else:
        print(f"❌ 分析结果文件不存在: {analysis_file}")
        return False
    
    # 2. 检查图表文件
    print("\n2. 检查图表文件...")
    if charts_dir.exists():
        chart_files = list(charts_dir.glob("*.png"))
        print(f"✅ 图表目录存在: {charts_dir}")
        print(f"📊 找到 {len(chart_files)} 个图表文件")
        
        # 显示前几个图表文件
        for i, chart in enumerate(chart_files[:5]):
            file_size = chart.stat().st_size
            print(f"  🖼️ {chart.name} ({file_size:,} 字节)")
    else:
        print(f"❌ 图表目录不存在: {charts_dir}")
        return False
    
    # 3. 直接测试Word报告生成
    print("\n3. 直接测试Word报告生成...")
    try:
        from backend.services.r_analysis import RAnalysisEngine
        
        engine = RAnalysisEngine()
        
        # 生成一个简化的DeepSeek报告用于测试
        mock_deepseek_report = """
# 压力采集数据质量分析报告

## 分析概述
本次分析对机器人压力采集系统进行了全面的质量评估。

## 主要发现
1. 系统整体表现良好
2. 各目标力值的精度符合要求
3. 建议继续监控过程稳定性

## 建议
- 维持当前校准参数
- 定期进行系统维护
        """
        
        print("📝 开始生成综合Word报告...")
        
        try:
            report_path = engine.generate_comprehensive_word_report(
                task_id=task_id,
                analysis_data=analysis_data,
                deepseek_report=mock_deepseek_report
            )
        except Exception as e:
            print(f"❌ Word报告生成异常: {e}")
            print("\n🔍 详细错误追踪:")
            tb = traceback.format_exc()
            print(tb)
            
            # 尝试找到具体的错误行
            lines = tb.split('\n')
            for i, line in enumerate(lines):
                if "'list' object has no attribute 'items'" in line:
                    print(f"\n🎯 错误位置: {line}")
                    if i > 0:
                        print(f"上一行: {lines[i-1]}")
                    if i < len(lines) - 1:
                        print(f"下一行: {lines[i+1]}")
                elif "File" in line and "line" in line:
                    print(f"📍 {line}")
            
            return False
        
        if report_path and Path(report_path).exists():
            file_size = Path(report_path).stat().st_size
            print(f"✅ Word报告生成成功!")
            print(f"📄 文件路径: {report_path}")
            print(f"📊 文件大小: {file_size:,} 字节")
            
            # 检查Word文件内容
            print("\n4. 检查Word文件内容...")
            try:
                from docx import Document
                doc = Document(report_path)
                
                print(f"📋 段落数量: {len(doc.paragraphs)}")
                print(f"📊 表格数量: {len(doc.tables)}")
                
                # 检查图片
                inline_shapes = []
                for paragraph in doc.paragraphs:
                    for run in paragraph.runs:
                        if run._element.xpath('.//pic:pic'):
                            inline_shapes.extend(run._element.xpath('.//pic:pic'))
                
                print(f"🖼️ 图片数量: {len(inline_shapes)}")
                
                if len(inline_shapes) > 0:
                    print("✅ Word文档包含图片!")
                else:
                    print("⚠️ Word文档中没有找到图片")
                
                # 显示部分段落内容
                print("\n📖 前几个段落内容:")
                for i, para in enumerate(doc.paragraphs[:5]):
                    if para.text.strip():
                        print(f"  {i+1}. {para.text[:100]}{'...' if len(para.text) > 100 else ''}")
                
                return True
                
            except Exception as e:
                print(f"❌ 检查Word文件失败: {e}")
                import traceback
                traceback.print_exc()
                return False
        else:
            print("❌ Word报告生成失败")
            return False
            
    except Exception as e:
        print(f"❌ Word报告生成异常: {e}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    success = debug_word_generation()
    
    if success:
        print("\n🎉 Word报告生成调试成功!")
    else:
        print("\n❌ Word报告生成调试失败")
        print("\n可能的解决方案:")
        print("• 检查所有文件路径")
        print("• 验证图表文件完整性")
        print("• 检查docx库是否正确安装")
        print("• 查看详细错误日志") 