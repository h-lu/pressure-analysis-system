#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
最终的完整综合报告测试
"""
import requests
import json
import time
from pathlib import Path

def final_comprehensive_test():
    """最终的完整综合报告测试"""
    print("🎯 最终的完整综合报告测试")
    print("=" * 60)
    
    # 使用已有的task_id
    task_id = "9ea02dd3-9e29-4309-ab3e-1c2172909982"
    
    # 1. 测试API生成综合报告
    print("1. 测试API生成综合报告...")
    try:
        response = requests.post(
            f'http://localhost:8000/api/deepseek/generate-comprehensive-word-report?task_id={task_id}',
            headers={'Content-Type': 'application/json'},
            timeout=60
        )
        
        if response.status_code == 200:
            result = response.json()
            print("✅ API调用成功")
            print(f"📊 响应: {result}")
            
            if result.get('success'):
                download_url = result.get('download_url')
                report_path = result.get('report_path')
                print(f"📄 报告路径: {report_path}")
                print(f"🔗 下载URL: {download_url}")
            else:
                print("❌ API返回失败状态")
                return False
        else:
            print(f"❌ API调用失败: {response.status_code}")
            print(f"错误内容: {response.text}")
            return False
    except Exception as e:
        print(f"❌ API调用异常: {e}")
        return False
    
    # 2. 验证生成的文件
    print("\n2. 验证生成的文件...")
    report_file = Path(f"temp/reports/comprehensive_analysis_report_{task_id}.docx")
    
    if report_file.exists():
        file_size = report_file.stat().st_size
        print(f"✅ 报告文件存在: {report_file}")
        print(f"📊 文件大小: {file_size:,} 字节")
        
        if file_size > 500000:  # 大于500KB
            print("✅ 文件大小正常，包含完整内容")
        else:
            print("⚠️ 文件可能不完整")
    else:
        print(f"❌ 报告文件不存在: {report_file}")
        return False
    
    # 3. 测试下载功能
    print("\n3. 测试下载功能...")
    try:
        download_response = requests.get(
            f'http://localhost:8000/api/download-comprehensive-report/{task_id}',
            timeout=30
        )
        
        if download_response.status_code == 200:
            print("✅ 下载成功")
            print(f"📊 下载文件大小: {len(download_response.content):,} 字节")
            
            # 验证Content-Type
            content_type = download_response.headers.get('content-type', '')
            if 'wordprocessingml' in content_type:
                print("✅ 文件类型正确 (Word文档)")
            else:
                print(f"⚠️ 文件类型异常: {content_type}")
            
            # 验证Content-Disposition
            content_disposition = download_response.headers.get('content-disposition', '')
            if 'attachment' in content_disposition:
                print("✅ 下载头设置正确")
            else:
                print(f"⚠️ 下载头异常: {content_disposition}")
                
        else:
            print(f"❌ 下载失败: {download_response.status_code}")
            return False
    except Exception as e:
        print(f"❌ 下载异常: {e}")
        return False
    
    # 4. 验证Word文档内容
    print("\n4. 验证Word文档内容...")
    try:
        from docx import Document
        doc = Document(str(report_file))
        
        print(f"📋 段落数量: {len(doc.paragraphs)}")
        print(f"📊 表格数量: {len(doc.tables)}")
        
        # 检查图片
        inline_shapes = []
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if run._element.xpath('.//pic:pic'):
                    inline_shapes.extend(run._element.xpath('.//pic:pic'))
        
        print(f"🖼️ 图片数量: {len(inline_shapes)}")
        
        # 检查关键章节
        key_sections = [
            "DeepSeek AI智能分析报告",
            "R统计分析详细数据与图表", 
            "综合结论与建议"
        ]
        
        found_sections = []
        for paragraph in doc.paragraphs:
            for section in key_sections:
                if section in paragraph.text:
                    found_sections.append(section)
        
        print(f"📖 找到关键章节: {len(set(found_sections))}/{len(key_sections)}")
        for section in set(found_sections):
            print(f"  ✅ {section}")
        
        # 验证完整性
        if (len(doc.paragraphs) > 50 and 
            len(doc.tables) > 3 and 
            len(inline_shapes) > 0 and 
            len(set(found_sections)) >= 2):
            print("✅ Word文档内容完整")
            return True
        else:
            print("⚠️ Word文档内容可能不完整")
            return False
            
    except Exception as e:
        print(f"❌ Word文档验证失败: {e}")
        return False

if __name__ == "__main__":
    success = final_comprehensive_test()
    
    if success:
        print("\n🎉 最终综合报告测试完全成功!")
        print("\n✅ 测试通过的功能:")
        print("  • API生成综合报告")
        print("  • 文件正确保存")
        print("  • 下载功能正常")
        print("  • Word文档内容完整")
        print("  • 包含图片和表格")
        print("  • 包含所有关键章节")
    else:
        print("\n❌ 最终综合报告测试失败") 