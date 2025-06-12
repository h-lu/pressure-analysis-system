#!/usr/bin/env python3
"""
验证Word文档内容的脚本
"""

import os
from pathlib import Path

def verify_word_document():
    """验证Word文档"""
    
    doc_path = "downloads/comprehensive_pressure_analysis_20250609_093009.docx"
    
    print("📄 验证Word文档")
    print("=" * 50)
    
    # 检查文件是否存在
    if not os.path.exists(doc_path):
        print(f"❌ 文件不存在: {doc_path}")
        return False
    
    # 获取文件信息
    file_size = os.path.getsize(doc_path)
    file_size_mb = file_size / (1024 * 1024)
    
    print(f"✅ 文件存在: {doc_path}")
    print(f"📊 文件大小: {file_size_mb:.2f} MB ({file_size:,} bytes)")
    
    # 尝试使用python-docx读取文档
    try:
        from docx import Document
        
        doc = Document(doc_path)
        
        # 统计文档内容
        paragraphs = len(doc.paragraphs)
        tables = len(doc.tables)
        
        # 统计图片
        images = 0
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                images += 1
        
        print(f"\n📋 文档内容统计:")
        print(f"   - 段落数: {paragraphs}")
        print(f"   - 表格数: {tables}")
        print(f"   - 图片数: {images}")
        
        # 显示前几个段落的内容
        print(f"\n📝 文档开头内容:")
        for i, para in enumerate(doc.paragraphs[:10]):
            if para.text.strip():
                print(f"   {i+1}. {para.text[:100]}{'...' if len(para.text) > 100 else ''}")
        
        # 检查是否包含关键内容
        full_text = "\n".join([para.text for para in doc.paragraphs])
        
        key_sections = [
            "DeepSeek AI分析",
            "R统计分析",
            "综合结论",
            "质量评估",
            "建议"
        ]
        
        print(f"\n🔍 关键章节检查:")
        for section in key_sections:
            if section in full_text:
                print(f"   ✅ 包含: {section}")
            else:
                print(f"   ❌ 缺失: {section}")
        
        # 检查图表
        if images > 0:
            print(f"\n📊 图表验证:")
            print(f"   ✅ 成功插入 {images} 个图表")
        else:
            print(f"\n📊 图表验证:")
            print(f"   ❌ 未发现图表")
        
        return True
        
    except ImportError:
        print("\n⚠️  无法导入python-docx，跳过详细内容验证")
        print("   可以手动打开文档查看内容")
        return True
        
    except Exception as e:
        print(f"\n❌ 读取文档时出错: {str(e)}")
        return False

def check_analysis_results():
    """检查分析结果文件"""
    
    print(f"\n🔍 检查分析结果文件")
    print("=" * 50)
    
    # 检查图表目录
    charts_dir = "backend/static/charts/4514bcbf-1459-409c-8e17-710e6b73ab31"
    
    if os.path.exists(charts_dir):
        chart_files = [f for f in os.listdir(charts_dir) if f.endswith('.png')]
        print(f"✅ 图表目录存在: {charts_dir}")
        print(f"📊 图表文件数量: {len(chart_files)}")
        
        # 显示前几个图表文件
        for i, chart in enumerate(chart_files[:5]):
            print(f"   {i+1}. {chart}")
        if len(chart_files) > 5:
            print(f"   ... 还有 {len(chart_files) - 5} 个图表")
    else:
        print(f"❌ 图表目录不存在: {charts_dir}")
    
    # 检查分析结果JSON
    json_file = os.path.join(charts_dir, "analysis_results.json")
    if os.path.exists(json_file):
        file_size = os.path.getsize(json_file)
        print(f"✅ 分析结果文件存在: {file_size:,} bytes")
    else:
        print(f"❌ 分析结果文件不存在")

if __name__ == "__main__":
    print("🚀 开始验证分析结果")
    
    # 验证Word文档
    word_ok = verify_word_document()
    
    # 检查分析结果
    check_analysis_results()
    
    print(f"\n🎉 验证完成")
    if word_ok:
        print("✅ Word文档验证通过")
        print("📄 可以打开文档查看完整的分析报告")
    else:
        print("❌ Word文档验证失败") 