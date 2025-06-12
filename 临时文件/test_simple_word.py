#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
简单的Word文档生成测试
"""
import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def test_simple_word_generation():
    """测试简单的Word文档生成"""
    print("🔬 测试简单Word文档生成...")
    
    try:
        # 创建Word文档
        doc = Document()
        
        # 设置文档样式
        style = doc.styles['Normal']
        style.font.name = 'SimSun'
        style.font.size = Pt(11)
        
        # 添加标题
        title = doc.add_heading('测试Word文档', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加内容
        doc.add_paragraph('这是一个测试段落。')
        doc.add_paragraph('测试中文字符显示。')
        
        # 创建目录
        reports_dir = Path("temp/reports")
        reports_dir.mkdir(parents=True, exist_ok=True)
        
        # 保存文档
        report_path = reports_dir / "test_simple_word.docx"
        doc.save(str(report_path))
        
        print(f"✅ 简单Word文档生成成功: {report_path}")
        
        # 检查文件
        if report_path.exists():
            file_size = os.path.getsize(report_path)
            print(f"✅ 文件大小: {file_size} 字节")
            return True
        else:
            print("❌ 文件未创建")
            return False
            
    except Exception as e:
        print(f"❌ Word文档生成失败: {e}")
        return False

if __name__ == "__main__":
    test_simple_word_generation() 