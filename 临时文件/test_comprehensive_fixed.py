#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
测试修复后的综合Word报告生成
"""
import sys
import os
import json
from pathlib import Path

# 添加backend到path
sys.path.append(os.path.join(os.path.dirname(__file__), 'backend'))

def test_comprehensive_fixed():
    """测试修复后的综合报告生成"""
    task_id = "9ea02dd3-9e29-4309-ab3e-1c2172909982"
    
    print(f"🔧 测试修复后的综合Word报告生成，task_id: {task_id}")
    print("=" * 60)
    
    try:
        from backend.services.r_analysis import RAnalysisEngine
        
        # 读取分析数据
        charts_dir = Path("backend/static/charts") / task_id
        analysis_file = charts_dir / "analysis_results.json"
        
        with open(analysis_file, 'r', encoding='utf-8') as f:
            analysis_data = json.load(f)
        
        print("✅ 分析数据读取成功")
        
        # 验证数据结构
        print("📊 验证数据结构:")
        print(f"  data_summary: {type(analysis_data.get('data_summary'))}")
        print(f"  overall_stats: {type(analysis_data.get('overall_stats'))}")
        print(f"  target_analysis: {type(analysis_data.get('target_analysis'))}")
        
        engine = RAnalysisEngine()
        
        # 生成一个简化的DeepSeek报告用于测试
        mock_deepseek_report = """
# 压力采集数据质量分析报告

## 分析概述
本次分析对机器人压力采集系统进行了全面的质量评估。

## 主要发现
1. 系统整体表现良好
2. 各目标力值的精度符合要求
3. 建议继续监控过程稳定性

## 建议
- 维持当前校准参数
- 定期进行系统维护
        """
        
        print("📝 开始生成综合Word报告...")
        
        # 直接调用方法，捕获详细错误
        try:
            report_path = engine.generate_comprehensive_word_report(
                task_id=task_id,
                analysis_data=analysis_data,
                deepseek_report=mock_deepseek_report
            )
            
            if report_path and Path(report_path).exists():
                file_size = Path(report_path).stat().st_size
                print(f"✅ Word报告生成成功!")
                print(f"📄 文件路径: {report_path}")
                print(f"📊 文件大小: {file_size:,} 字节")
                
                # 检查Word文件内容
                print("\n📖 检查Word文件内容...")
                try:
                    from docx import Document
                    doc = Document(report_path)
                    
                    print(f"📋 段落数量: {len(doc.paragraphs)}")
                    print(f"📊 表格数量: {len(doc.tables)}")
                    
                    # 检查图片
                    inline_shapes = []
                    for paragraph in doc.paragraphs:
                        for run in paragraph.runs:
                            if run._element.xpath('.//pic:pic'):
                                inline_shapes.extend(run._element.xpath('.//pic:pic'))
                    
                    print(f"🖼️ 图片数量: {len(inline_shapes)}")
                    
                    if len(inline_shapes) > 0:
                        print("✅ Word文档包含图片!")
                    else:
                        print("⚠️ Word文档中没有找到图片")
                    
                    return True
                    
                except Exception as e:
                    print(f"❌ 检查Word文件失败: {e}")
                    return False
            else:
                print("❌ Word报告生成失败")
                return False
                
        except Exception as e:
            print(f"❌ Word报告生成异常: {e}")
            
            # 详细错误追踪
            import traceback
            tb = traceback.format_exc()
            print("\n🔍 详细错误追踪:")
            print(tb)
            
            # 查找具体错误位置
            lines = tb.split('\n')
            for i, line in enumerate(lines):
                if "'list' object has no attribute 'items'" in line:
                    print(f"\n🎯 错误位置: {line}")
                    # 查找相关的文件和行号
                    for j in range(max(0, i-10), min(len(lines), i+3)):
                        if "File" in lines[j] and "line" in lines[j]:
                            print(f"📍 {lines[j]}")
                        elif lines[j].strip() and not lines[j].startswith("Traceback"):
                            print(f"   {lines[j]}")
            
            return False
        
    except Exception as e:
        print(f"❌ 测试失败: {e}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    success = test_comprehensive_fixed()
    
    if success:
        print("\n🎉 综合Word报告测试成功!")
    else:
        print("\n❌ 综合Word报告测试失败") 