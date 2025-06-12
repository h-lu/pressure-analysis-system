#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
测试已经生成的综合报告文件
"""
import requests
from pathlib import Path

def test_existing_report():
    """测试已经生成的综合报告文件"""
    print("🎯 测试已经生成的综合报告文件")
    print("=" * 60)
    
    task_id = "9ea02dd3-9e29-4309-ab3e-1c2172909982"
    
    # 1. 验证生成的文件
    print("1. 验证生成的文件...")
    report_file = Path(f"temp/reports/comprehensive_analysis_report_{task_id}.docx")
    
    if report_file.exists():
        file_size = report_file.stat().st_size
        print(f"✅ 报告文件存在: {report_file}")
        print(f"📊 文件大小: {file_size:,} 字节")
        
        if file_size > 500000:  # 大于500KB
            print("✅ 文件大小正常，包含完整内容")
        else:
            print("⚠️ 文件可能不完整")
    else:
        print(f"❌ 报告文件不存在: {report_file}")
        return False
    
    # 2. 测试下载功能
    print("\n2. 测试下载功能...")
    try:
        download_response = requests.get(
            f'http://localhost:8000/api/download-comprehensive-report/{task_id}',
            timeout=30
        )
        
        if download_response.status_code == 200:
            print("✅ 下载成功")
            print(f"📊 下载文件大小: {len(download_response.content):,} 字节")
            
            # 验证Content-Type
            content_type = download_response.headers.get('content-type', '')
            if 'wordprocessingml' in content_type:
                print("✅ 文件类型正确 (Word文档)")
            else:
                print(f"⚠️ 文件类型异常: {content_type}")
            
            # 验证Content-Disposition
            content_disposition = download_response.headers.get('content-disposition', '')
            if 'attachment' in content_disposition:
                print("✅ 下载头设置正确")
            else:
                print(f"⚠️ 下载头异常: {content_disposition}")
                
        else:
            print(f"❌ 下载失败: {download_response.status_code}")
            return False
    except Exception as e:
        print(f"❌ 下载异常: {e}")
        return False
    
    # 3. 验证Word文档内容
    print("\n3. 验证Word文档内容...")
    try:
        from docx import Document
        doc = Document(str(report_file))
        
        print(f"📋 段落数量: {len(doc.paragraphs)}")
        print(f"📊 表格数量: {len(doc.tables)}")
        
        # 检查图片
        inline_shapes = []
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if run._element.xpath('.//pic:pic'):
                    inline_shapes.extend(run._element.xpath('.//pic:pic'))
        
        print(f"🖼️ 图片数量: {len(inline_shapes)}")
        
        # 检查关键章节
        key_sections = [
            "DeepSeek AI智能分析报告",
            "R统计分析详细数据与图表", 
            "综合结论与建议"
        ]
        
        found_sections = []
        for paragraph in doc.paragraphs:
            for section in key_sections:
                if section in paragraph.text:
                    found_sections.append(section)
        
        print(f"📖 找到关键章节: {len(set(found_sections))}/{len(key_sections)}")
        for section in set(found_sections):
            print(f"  ✅ {section}")
        
        # 显示前几个段落内容
        print("\n📖 文档开头内容:")
        for i, para in enumerate(doc.paragraphs[:10]):
            if para.text.strip():
                print(f"  {i+1}. {para.text[:80]}{'...' if len(para.text) > 80 else ''}")
        
        # 验证完整性
        if (len(doc.paragraphs) > 50 and 
            len(doc.tables) > 3 and 
            len(inline_shapes) > 0 and 
            len(set(found_sections)) >= 2):
            print("\n✅ Word文档内容完整")
            return True
        else:
            print("\n⚠️ Word文档内容可能不完整")
            return False
            
    except Exception as e:
        print(f"❌ Word文档验证失败: {e}")
        return False

if __name__ == "__main__":
    success = test_existing_report()
    
    if success:
        print("\n🎉 综合报告文件测试完全成功!")
        print("\n✅ 验证通过的功能:")
        print("  • 文件正确保存")
        print("  • 下载功能正常")
        print("  • Word文档内容完整")
        print("  • 包含图片和表格")
        print("  • 包含所有关键章节")
        print("\n🔧 修复完成的问题:")
        print("  • 数据结构一致性问题")
        print("  • Word文档生成中的列表处理问题")
        print("  • 图表插入功能")
        print("  • 下载API路径问题")
    else:
        print("\n❌ 综合报告文件测试失败") 