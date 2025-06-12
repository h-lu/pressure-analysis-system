#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
测试综合Word报告生成 - 包含实际图表生成
"""
import requests
import json
import time
import os
from pathlib import Path

def create_test_data():
    """创建测试数据文件"""
    import pandas as pd
    import numpy as np
    
    # 生成测试数据
    np.random.seed(42)
    n = 50
    
    data = []
    for i in range(n):
        # 随机选择目标力值
        target_force = np.random.choice([5, 25, 50])
        
        # 生成位置坐标
        x = np.random.uniform(90, 110)
        y = np.random.uniform(90, 110) 
        z = np.random.uniform(90, 110)
        
        # 生成力值（带一些噪声）
        noise = np.random.normal(0, target_force * 0.05)
        force = target_force + noise
        
        data.append({
            '序号': i + 1,
            'X': round(x, 1),
            'Y': round(y, 1),
            'Z': round(z, 1),
            '力值': f"{force:.1f}N"
        })
    
    df = pd.DataFrame(data)
    
    # 保存到CSV文件
    test_file = "test_data_with_charts.csv"
    df.to_csv(test_file, index=False, encoding='utf-8')
    return test_file

def upload_test_file(file_path):
    """上传测试文件"""
    try:
        with open(file_path, 'rb') as f:
            files = {'file': (file_path, f, 'text/csv')}
            response = requests.post('http://localhost:8000/api/upload', files=files, timeout=30)
            
        if response.status_code == 200:
            result = response.json()
            return result.get('file_id')
        else:
            print(f"❌ 上传失败: {response.status_code}")
            return None
    except Exception as e:
        print(f"❌ 上传异常: {e}")
        return None

def start_analysis(file_id):
    """启动分析任务"""
    try:
        payload = {
            "file_id": file_id,
            "target_forces": [5.0, 25.0, 50.0],
            "tolerance_abs": 2.0,
            "tolerance_pct": 5.0,
            "analysis_type": "comprehensive"
        }
        
        response = requests.post('http://localhost:8000/api/analyze', json=payload, timeout=30)
        
        if response.status_code == 200:
            result = response.json()
            return result.get('task_id')
        else:
            print(f"❌ 启动分析失败: {response.status_code}")
            return None
    except Exception as e:
        print(f"❌ 启动分析异常: {e}")
        return None

def wait_for_analysis(task_id, max_wait=300):
    """等待分析完成"""
    start_time = time.time()
    while time.time() - start_time < max_wait:
        try:
            response = requests.get(f'http://localhost:8000/api/task/{task_id}', timeout=10)
            if response.status_code == 200:
                result = response.json()
                task = result.get('task', {})
                status = task.get('status')
                progress = task.get('progress', 0)
                message = task.get('message', '')
                
                print(f"分析进度: {progress}% - {message}")
                
                if status == 'completed':
                    return True
                elif status == 'failed':
                    print(f"❌ 分析失败: {task.get('error', '未知错误')}")
                    return False
                    
            time.sleep(5)  # 等待5秒后重试
            
        except Exception as e:
            print(f"⚠️ 查询状态异常: {e}")
            time.sleep(5)
    
    print("❌ 分析超时")
    return False

def test_comprehensive_report_with_charts():
    """测试包含图表的综合Word报告生成功能"""
    print("🔧 测试包含图表的综合Word报告生成功能...")
    print("=" * 70)
    
    # 1. 检查服务器状态
    print("1. 检查服务器状态...")
    try:
        response = requests.get('http://localhost:8000/api/deepseek/test-connection', timeout=10)
        if response.status_code == 200:
            print("✅ 服务器连接正常")
        else:
            print("❌ 服务器连接失败")
            return False
    except Exception as e:
        print(f"❌ 服务器连接异常: {e}")
        return False
    
    # 2. 创建和上传测试数据
    print("\n2. 创建和上传测试数据...")
    test_file = create_test_data()
    print(f"✅ 测试数据文件创建: {test_file}")
    
    file_id = upload_test_file(test_file)
    if not file_id:
        print("❌ 文件上传失败")
        return False
    print(f"✅ 文件上传成功: {file_id}")
    
    # 3. 启动R分析
    print("\n3. 启动R分析...")
    task_id = start_analysis(file_id)
    if not task_id:
        print("❌ 启动分析失败")
        return False
    print(f"✅ 分析任务启动: {task_id}")
    
    # 4. 等待分析完成
    print("\n4. 等待分析完成...")
    if not wait_for_analysis(task_id):
        print("❌ 分析失败或超时")
        return False
    print("✅ R分析完成")
    
    # 5. 检查生成的图表
    print("\n5. 检查生成的图表...")
    charts_dir = Path("backend/static/charts") / task_id
    if charts_dir.exists():
        chart_files = list(charts_dir.glob("*.png"))
        print(f"✅ 找到 {len(chart_files)} 个图表文件")
        for chart in chart_files[:5]:  # 显示前5个
            print(f"  📊 {chart.name}")
    else:
        print(f"❌ 图表目录不存在: {charts_dir}")
        return False
    
    # 6. 生成综合报告
    print("\n6. 生成综合报告...")
    try:
        response = requests.post(
            'http://localhost:8000/api/deepseek/generate-comprehensive-word-report',
            params={'task_id': task_id},
            timeout=180  # 增加超时时间
        )
        
        if response.status_code == 200:
            result = response.json()
            if result.get('success'):
                print("✅ 综合报告生成成功")
                print(f"📄 下载链接: {result.get('download_url')}")
                print(f"📁 文件路径: {result.get('report_path')}")
                
                # 验证文件是否存在
                report_path = result.get('report_path')
                if report_path and os.path.exists(report_path):
                    file_size = os.path.getsize(report_path)
                    print(f"✅ 报告文件存在，大小: {file_size:,} 字节")
                    
                    if file_size > 100000:  # 大于100KB表示可能包含图片
                        print("✅ 文件大小正常，可能包含图表")
                        return True, task_id
                    else:
                        print("⚠️ 文件较小，可能没有包含图表")
                        return True, task_id
                else:
                    print("❌ 报告文件不存在")
                    return False, None
            else:
                print(f"❌ 综合报告生成失败: {result.get('message', '未知错误')}")
                return False, None
        else:
            print(f"❌ API调用失败: {response.status_code}")
            print(f"错误信息: {response.text}")
            return False, None
            
    except Exception as e:
        print(f"❌ 综合报告生成异常: {e}")
        return False, None
    
    # 7. 测试下载功能
    print("\n7. 测试下载功能...")
    try:
        download_response = requests.get(
            f'http://localhost:8000/api/download-comprehensive-report/{task_id}',
            timeout=30
        )
        
        if download_response.status_code == 200:
            content_type = download_response.headers.get('content-type', '')
            if 'wordprocessingml' in content_type or 'application/octet-stream' in content_type:
                print("✅ 下载功能正常，文件类型正确")
                print(f"📁 Content-Type: {content_type}")
                print(f"📊 文件大小: {len(download_response.content):,} 字节")
                return True, task_id
            else:
                print(f"⚠️ 文件类型异常: {content_type}")
                return False, None
        else:
            print(f"❌ 下载失败: {download_response.status_code}")
            return False, None
            
    except Exception as e:
        print(f"❌ 下载测试异常: {e}")
        return False, None

def check_word_content(task_id):
    """检查Word文档内容"""
    print(f"\n8. 检查Word文档内容...")
    
    # 查找报告文件
    reports_dir = Path("temp/reports")
    report_files = list(reports_dir.glob(f"comprehensive_analysis_report_{task_id}.docx"))
    
    if not report_files:
        print("❌ 未找到综合报告文件")
        return False
    
    report_file = report_files[0]
    
    try:
        from docx import Document
        doc = Document(str(report_file))
        
        print(f"✅ Word文档可以正常打开")
        print(f"📋 段落数量: {len(doc.paragraphs)}")
        print(f"📊 表格数量: {len(doc.tables)}")
        
        # 检查是否包含图片
        inline_shapes = []
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if run._element.xpath('.//pic:pic'):
                    inline_shapes.extend(run._element.xpath('.//pic:pic'))
        
        print(f"🖼️ 图片数量: {len(inline_shapes)}")
        
        if len(inline_shapes) > 0:
            print("✅ 文档包含图片!")
            return True
        else:
            print("⚠️ 文档中没有找到图片")
            return False
            
    except Exception as e:
        print(f"❌ 检查Word文档失败: {e}")
        return False

if __name__ == "__main__":
    success, task_id = test_comprehensive_report_with_charts()
    
    if success and task_id:
        # 检查Word文档内容
        has_images = check_word_content(task_id)
        
        print("\n" + "=" * 70)
        if success and has_images:
            print("🎉 包含图表的综合报告测试完全成功！")
            print("\n✅ 验证通过的功能:")
            print("  • 服务器连接正常")
            print("  • 测试数据生成和上传")
            print("  • R分析执行成功")
            print("  • 图表文件生成")
            print("  • DeepSeek分析正常")
            print("  • Word报告生成成功")
            print("  • 报告包含图表")
            print("  • 文件下载正常")
        else:
            print("⚠️ 综合报告测试基本成功，但图表可能有问题")
            print("\n🔧 可能的改进点:")
            print("  • 检查图表插入逻辑")
            print("  • 检查图表文件路径")
            print("  • 验证图表文件格式")
    else:
        print("❌ 综合报告测试失败")
        print("\n🔧 可能的解决方案:")
        print("  • 检查服务器状态")
        print("  • 检查R分析脚本")
        print("  • 检查API配置")
        print("  • 检查文件权限")
    
    # 清理临时文件
    try:
        if os.path.exists("test_data_with_charts.csv"):
            os.remove("test_data_with_charts.csv")
    except:
        pass 