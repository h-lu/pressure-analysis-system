#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
最简单的Word报告测试
"""
import sys
import os
import json
from pathlib import Path

# 添加backend到path
sys.path.append(os.path.join(os.path.dirname(__file__), 'backend'))

def simple_word_test():
    """最简单的Word报告测试"""
    task_id = "9ea02dd3-9e29-4309-ab3e-1c2172909982"
    
    print(f"🔧 简单Word报告测试，task_id: {task_id}")
    print("=" * 60)
    
    try:
        from backend.services.r_analysis import RAnalysisEngine
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from datetime import datetime
        
        # 读取分析数据
        charts_dir = Path("backend/static/charts") / task_id
        analysis_file = charts_dir / "analysis_results.json"
        
        with open(analysis_file, 'r', encoding='utf-8') as f:
            analysis_data = json.load(f)
        
        print("✅ 分析数据读取成功")
        
        # 创建最简单的Word文档
        doc = Document()
        
        # 设置文档样式
        style = doc.styles['Normal']
        style.font.name = 'SimSun'
        style.font.size = Pt(11)
        
        print("✅ Word文档创建成功")
        
        # 添加标题
        title = doc.add_heading('压力系统数据分析综合报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加生成时间
        doc.add_paragraph(f'报告生成时间：{datetime.now().strftime("%Y年%m月%d日 %H:%M:%S")}')
        doc.add_paragraph(f'任务ID：{task_id}')
        doc.add_paragraph('')
        
        print("✅ 基本信息添加成功")
        
        # 测试数据概览部分
        print("🔧 测试数据概览部分...")
        doc.add_heading('数据概览', level=1)
        data_summary = analysis_data.get('data_summary', {})
        
        if isinstance(data_summary, dict):
            print("✅ data_summary是字典格式")
            summary_table = doc.add_table(rows=1, cols=2)
            summary_table.style = 'Table Grid'
            
            hdr_cells = summary_table.rows[0].cells
            hdr_cells[0].text = '指标'
            hdr_cells[1].text = '数值'
            
            for key, value in data_summary.items():
                row_cells = summary_table.add_row().cells
                row_cells[0].text = str(key)
                row_cells[1].text = str(value)
            
            print("✅ 数据概览表格添加成功")
        else:
            print(f"❌ data_summary不是字典格式: {type(data_summary)}")
            return False
        
        # 测试整体统计部分
        print("🔧 测试整体统计部分...")
        doc.add_heading('整体统计', level=1)
        overall_stats = analysis_data.get('overall_stats', {})
        
        if isinstance(overall_stats, dict):
            print("✅ overall_stats是字典格式")
            stats_table = doc.add_table(rows=1, cols=2)
            stats_table.style = 'Table Grid'
            
            hdr_cells = stats_table.rows[0].cells
            hdr_cells[0].text = '统计指标'
            hdr_cells[1].text = '数值'
            
            for key, value in overall_stats.items():
                row_cells = stats_table.add_row().cells
                row_cells[0].text = str(key)
                if isinstance(value, (int, float)):
                    row_cells[1].text = f"{value:.3f}"
                else:
                    row_cells[1].text = str(value)
            
            print("✅ 整体统计表格添加成功")
        else:
            print(f"❌ overall_stats不是字典格式: {type(overall_stats)}")
            return False
        
        # 测试图表插入
        print("🔧 测试图表插入...")
        doc.add_heading('关键图表', level=1)
        
        # 查找一个图表文件
        chart_files = list(charts_dir.glob("*.png"))
        if chart_files:
            chart_file = chart_files[0]
            print(f"📊 找到图表文件: {chart_file.name}")
            
            try:
                doc.add_paragraph(f"图表：{chart_file.name}")
                doc.add_picture(str(chart_file), width=Inches(5.0))
                print("✅ 图表插入成功")
            except Exception as e:
                print(f"❌ 图表插入失败: {e}")
                return False
        else:
            print("⚠️ 未找到图表文件")
        
        # 保存文档
        output_path = f"temp/reports/simple_test_report_{task_id}.docx"
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)
        
        file_size = Path(output_path).stat().st_size
        print(f"✅ Word文档保存成功!")
        print(f"📄 文件路径: {output_path}")
        print(f"📊 文件大小: {file_size:,} 字节")
        
        return True
        
    except Exception as e:
        print(f"❌ 测试失败: {e}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    success = simple_word_test()
    
    if success:
        print("\n🎉 简单Word报告测试成功!")
    else:
        print("\n❌ 简单Word报告测试失败") 