#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
检查生成的Word文件内容
"""
import os
from pathlib import Path
from docx import Document

def check_word_file():
    """检查最新生成的Word文件"""
    reports_dir = Path("temp/reports")
    
    # 查找最新的综合报告文件
    word_files = list(reports_dir.glob("comprehensive_analysis_report_*.docx"))
    if not word_files:
        print("❌ 未找到综合报告文件")
        return
    
    # 按修改时间排序，取最新的
    latest_file = max(word_files, key=os.path.getmtime)
    print(f"📄 检查文件: {latest_file}")
    print(f"📊 文件大小: {latest_file.stat().st_size:,} 字节")
    
    try:
        # 尝试打开Word文档
        doc = Document(str(latest_file))
        
        print(f"✅ Word文档可以正常打开")
        print(f"📋 段落数量: {len(doc.paragraphs)}")
        print(f"📊 表格数量: {len(doc.tables)}")
        
        # 显示前几个段落的内容
        print("\n📖 文档开头内容:")
        print("-" * 50)
        for i, para in enumerate(doc.paragraphs[:10]):
            if para.text.strip():
                print(f"{i+1:2d}. {para.text.strip()}")
        print("-" * 50)
        
        # 检查是否包含关键章节
        all_text = '\n'.join([para.text for para in doc.paragraphs])
        
        key_sections = [
            "综合压力系统数据分析报告",
            "DeepSeek AI智能分析报告",
            "R统计分析详细数据与图表",
            "综合结论与建议"
        ]
        
        print("\n🔍 关键章节检查:")
        for section in key_sections:
            if section in all_text:
                print(f"✅ 包含: {section}")
            else:
                print(f"❌ 缺失: {section}")
        
        # 检查表格内容
        if doc.tables:
            print(f"\n📊 表格检查:")
            for i, table in enumerate(doc.tables[:3]):  # 检查前3个表格
                print(f"表格 {i+1}: {len(table.rows)} 行 x {len(table.columns)} 列")
                if table.rows:
                    first_row = [cell.text.strip() for cell in table.rows[0].cells]
                    print(f"  表头: {first_row}")
        
        print(f"\n✅ Word文件结构正常")
        return True
        
    except Exception as e:
        print(f"❌ Word文档打开失败: {e}")
        return False

if __name__ == "__main__":
    check_word_file() 