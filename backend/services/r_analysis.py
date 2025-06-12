"""
R分析引擎服务 - 通过subprocess调用R脚本进行数据分析
"""
import os
import logging
import subprocess
import json
from pathlib import Path
from typing import Dict, List, Any, Optional
import pandas as pd
from datetime import datetime
import pytz
from docx import Document

# 设置时区为上海时区
SHANGHAI_TZ = pytz.timezone('Asia/Shanghai')
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import shutil
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import RGBColor

from ..core.config import settings
from ..models.schemas import AnalysisParams

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class RAnalysisEngine:
    """R分析引擎，负责调用R脚本并处理结果"""
    
    def __init__(self):
        # 确保R的可执行文件路径在环境变量中
        if settings.R_HOME not in os.environ['PATH']:
            os.environ['PATH'] = f"{settings.R_HOME}{os.pathsep}{os.environ['PATH']}"
        
        # 修正R脚本的路径，使其指向 backend/r_analysis/
        self.r_script_path = os.path.join(settings.BASE_DIR, "backend", "r_analysis", "pressure_analysis.R")
        if not os.path.exists(self.r_script_path):
            raise FileNotFoundError(f"R脚本未找到: {self.r_script_path}")
            
        try:
            # 设置静态文件目录
            self.static_dir = Path(settings.STATIC_DIR)
            self.charts_dir = Path(settings.CHARTS_DIR)
            
            # 检查R是否可用
            result = subprocess.run(['R', '--version'], capture_output=True, text=True, timeout=10)
            if result.returncode != 0:
                raise Exception("R未安装或不可用")
            
            logger.info("R分析引擎初始化成功")
            
        except subprocess.TimeoutExpired:
            logger.error("R版本检查超时")
            raise Exception("R响应超时")
        except FileNotFoundError:
            logger.error("R未找到，请确保R已安装并在PATH中")
            raise Exception("R未安装")
        except Exception as e:
            logger.error(f"R分析引擎初始化失败: {str(e)}")
            raise
    
    def analyze_data(self, csv_path: str, params: AnalysisParams, task_id: str) -> Dict[str, Any]:
        """
        核心分析函数：调用R脚本执行数据分析。
        """
        # 1. 为每个任务创建一个独立的、带时间戳的输出目录
        output_dir = Path(settings.CHARTS_DIR) / task_id
        output_dir.mkdir(parents=True, exist_ok=True)
        
        # 2. 处理容差参数，确保它们是列表
        num_targets = len(params.target_forces)
        
        tolerance_abs_list = params.tolerance_abs
        if isinstance(tolerance_abs_list, float):
            tolerance_abs_list = [tolerance_abs_list] * num_targets
            
        tolerance_pct_list = params.tolerance_pct
        if isinstance(tolerance_pct_list, float):
            tolerance_pct_list = [tolerance_pct_list] * num_targets

        # 3. 构建R脚本的命令行参数
        cmd = [
            "Rscript",
            self.r_script_path,
            "--input", str(csv_path),
            "--output-dir", str(output_dir),
            "--file-id", task_id,
            "--target-forces", ",".join(map(str, params.target_forces)),
            "--tolerance-abs", ",".join(map(str, tolerance_abs_list)),
            "--tolerance-pct", ",".join(map(str, tolerance_pct_list)),
        ]
        
        logger.info(f"即将执行R命令: {' '.join(cmd)}")
        
        try:
            # 4. 执行R脚本，并捕获输出
            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                check=True,
                encoding='utf-8'
            )
            
            # 记录R脚本的输出，便于调试
            if result.stdout:
                logger.info("R script stdout:\n" + result.stdout)
            if result.stderr:
                logger.warning("R script stderr:\n" + result.stderr)

        except subprocess.CalledProcessError as e:
            # 如果R脚本执行失败，构造详细的错误信息
            error_message = f"R分析执行失败: {e.stderr}"
            logger.error(error_message, exc_info=True)
            raise Exception(error_message)
        
        # 5. 读取R脚本生成的JSON结果文件
        result_json_path = output_dir / "analysis_results.json"
        if not result_json_path.exists():
            error_msg = "R脚本执行成功，但未找到预期的结果文件 analysis_results.json"
            logger.error(error_msg)
            raise FileNotFoundError(error_msg)

        with open(result_json_path, 'r', encoding='utf-8') as f:
            analysis_results = json.load(f)

        # 6. 保存到历史记录
        self.save_to_history(task_id, analysis_results, params.file_id)
        
        return analysis_results
    
    def _process_results(self, task_id: str, output_dir: Path) -> Dict[str, Any]:
        """处理R脚本的输出结果"""
        results_file = output_dir / "analysis_results.json"
        if not results_file.exists():
            raise FileNotFoundError("分析结果文件 analysis_results.json 未找到")

        with open(results_file, 'r', encoding='utf-8') as f:
            statistics = json.load(f)

        def _to_dict(data: Any) -> Dict:
            """将列表中的单个对象转换为对象"""
            if isinstance(data, list) and len(data) > 0 and isinstance(data[0], dict):
                return data[0]
            if isinstance(data, dict):
                return data
            return {}

        # 确保关键数据是对象而非列表
        data_summary = _to_dict(statistics.get('data_summary'))
        overall_stats = _to_dict(statistics.get('overall_stats'))
        target_analysis = statistics.get('target_analysis', [])
        process_capability = statistics.get('process_capability', [])

        charts = self._collect_charts(task_id, output_dir)
        files = self._collect_files(task_id, output_dir)
        
        # 生成详细分析解读
        analysis_interpretation = self._generate_analysis_interpretation(statistics)
        
        logger.info(f"生成 {len(charts)} 个图表, {len(files)} 个文件")
        
        return {
            'task_id': task_id,
            'status': 'completed',
            'charts': charts,
            'files': files,
            'summary_report': self._generate_summary_report(statistics),
            'data_quality_report': self._generate_quality_report(statistics),
            'recommendations': self._generate_recommendations(statistics),
            'analysis_interpretation': analysis_interpretation,
            'data_summary': data_summary,
            'overall_stats': overall_stats,
            'target_analysis': target_analysis,
            'trend_stats': statistics.get('trend_stats', []),
            'outlier_summary': statistics.get('outlier_summary', []),
            'stability_analysis': statistics.get('stability_analysis', []),
            'change_point_analysis': statistics.get('change_point_analysis', []),
            'autocorr_analysis': statistics.get('autocorr_analysis', []),
            'process_capability': process_capability,
            
            # 新增高级分析结果
            'spatial_analysis': statistics.get('spatial_analysis', {}),
            'error_distribution_analysis': statistics.get('error_distribution_analysis', {}),
            'multi_source_variation_analysis': statistics.get('multi_source_variation_analysis', {}),
            
            'summary': statistics.get('summary', {}),
            'report_url': f'/api/download-report/{task_id}',
            'scatter_data_url': f'/static/charts/{task_id}/scatter_3d_data.csv'
        }
    
    def _collect_charts(self, task_id: str, output_dir: Path) -> List[Dict[str, Any]]:
        """收集生成的图表文件 - 完整的30个图表"""
        charts = []
        
        if not output_dir.exists():
            return charts
        
        # 完整的图表信息映射 - 包含所有30个图表
        chart_mapping = {
            # 基础分析图表 (1-5)
            "force_time_series.png": {
                "title": "力值时间序列图（按目标力值分组）",
                "category": "基础分析",
                "description": "展示力值随时间的变化趋势和波动模式，实心点为合格数据，空心点为异常数据",
                "interpretation": """
• 实心点：在容差内的合格数据点
• 空心点：超出容差的异常数据点  
• 红虚线：目标力值水平线
• 阴影区域：绝对容差范围
关注要点：
- 数据点是否主要集中在容差区域内
- 是否存在明显的上升或下降趋势
- 异常点是否存在聚集现象或周期性模式
                """
            },
            "force_histogram.png": {
                "title": "力值分布直方图（按目标力值分组）",
                "category": "基础分析", 
                "description": "分析每个目标力值组的数据分布特征和正态性",
                "interpretation": """
• 直方图：力值的频次分布
• 红虚线：目标力值位置
• 蓝实线：实际测量的平均值
关注要点：
- 分布是否接近正态分布（钟形曲线）
- 实际均值与目标值的偏离程度
- 分布的宽度（反映数据稳定性）
- 是否存在双峰或多峰分布
                """
            },
            "force_boxplot.png": {
                "title": "力值箱线图（按目标力值分组）",
                "category": "基础分析",
                "description": "快速识别数据的分位数特征和异常值分布",
                "interpretation": """
• 箱体：25%-75%分位数范围（IQR）
• 中线：中位数
• 白菱形：平均值
• 红点：统计异常值
• 红虚线：目标力值
关注要点：
- 箱体的高度（数据离散程度）
- 中位数与目标值的对齐程度
- 异常值的数量和分布
                """
            },
            "deviation_analysis.png": {
                "title": "绝对偏差箱线图（按目标力值分组）",
                "category": "偏差分析",
                "description": "分析实际测量值与目标值的绝对偏差分布",
                "interpretation": """
• Y轴：实际力值 - 目标力值
• 绿线：零偏差（理想状态）
• 橙虚线：绝对容差限制
• 正值：测量值高于目标
• 负值：测量值低于目标
关注要点：
- 偏差分布是否以零为中心
- 是否存在系统性偏移
- 超出容差限制的数据点比例
                """
            },
            "percentage_deviation.png": {
                "title": "百分比偏差箱线图（按目标力值分组）",
                "category": "偏差分析",
                "description": "分析相对于目标值的百分比偏差，消除目标值大小的影响",
                "interpretation": """
• Y轴：(实际力值-目标力值)/目标力值 × 100%
• 绿线：零偏差
• 橙虚线：百分比容差限制
关注要点：
- 不同目标力值组的相对精度是否一致
- 小力值和大力值的相对稳定性对比
- 百分比容差的实际达成情况
                """
            },
            
            # 统计过程控制图表 (6-12)
            "shewhart_control.png": {
                "title": "Shewhart控制图（按目标力值分组）",
                "category": "统计过程控制",
                "description": "统计过程控制的核心工具，监测过程稳定性",
                "interpretation": """
• 绿线：过程中心线（均值）
• 红虚线：3σ控制限（99.7%数据应在此范围内）
• 橙点线：2σ警戒线（95%数据应在此范围内）
• X标记：检测到的异常值
关注要点：
- 超出3σ控制限的点（过程失控）
- 连续7点在中心线同一侧（过程偏移）
- 连续趋势或周期性模式
                """
            },
            "moving_average.png": {
                "title": "移动平均图（按目标力值分组）",
                "category": "统计过程控制",
                "description": "平滑短期波动，突出长期趋势变化",
                "interpretation": """
• 细线：原始数据
• 粗线：移动平均线
• 彩色带：移动标准差范围
• 红虚线：目标力值
关注要点：
- 移动平均线是否稳定在目标值附近
- 标准差带的宽度变化（稳定性变化）
- 长期趋势的方向和幅度
                """
            },
            "cusum_chart.png": {
                "title": "CUSUM累计和控制图",
                "category": "统计过程控制",
                "description": "检测过程均值的小幅持续性偏移，对微小变化敏感",
                "interpretation": """
• 红线：累计正偏差（高于目标的累积）
• 蓝线：累计负偏差（低于目标的累积）
• 水平虚线：±4的决策界限
关注要点：
- 超出±4界限表示过程失控
- 持续上升/下降趋势表示系统偏移
- 比Shewhart图对小偏移更敏感
                """
            },
            "ewma_chart.png": {
                "title": "EWMA指数加权移动平均控制图",
                "category": "统计过程控制",
                "description": "对小幅偏移敏感的统计控制，具有记忆功能",
                "interpretation": """
• 粗线：EWMA曲线（平滑的指数加权移动平均）
• 绿线：目标中心线
• 红虚线：3σ控制限
• 蓝色阴影：控制区间
关注要点：
- 平滑参数λ=0.2，对历史数据有记忆
- 比传统控制图更快检测小偏移
- EWMA线的连续趋势比单点更重要
                """
            },
            "imr_chart.png": {
                "title": "I-MR个值移动极差控制图",
                "category": "统计过程控制",
                "description": "适用于个别测量值的统计过程控制",
                "interpretation": """
• I图(上)：个别测量值的控制
• MR图(下)：相邻测量值间变异的控制
• 绿线：过程中心线
• 红虚线：控制限
关注要点：
- I图失控表示过程位置偏移
- MR图失控表示过程变异增大
- 适用于单个测量值的连续监控
                """
            },
            "xbar_r_chart.png": {
                "title": "X-bar & R控制图组合",
                "category": "统计过程控制",
                "description": "同时监控过程均值和变异性的组合控制图",
                "interpretation": """
• X-bar图：监控过程平均水平
• R图：监控过程变异性（极差）
• 绿线：中心线
• 红虚线：控制限
关注要点：
- X-bar图失控表示均值偏移
- R图失控表示变异增大
- 两图需结合分析，识别不同类型的过程变化
                """
            },
            "run_chart.png": {
                "title": "运行图（Run Chart）",
                "category": "统计过程控制",
                "description": "检测数据中的非随机模式和趋势",
                "interpretation": """
• 蓝虚线：中位数
• 三角形：高于/低于中位数的点
• 连续游程：连续的同向偏离
关注要点：
- 连续8点在中位数同一侧（非随机模式）
- 游程长度的分布
- 明显的趋势或周期性
                """
            },
            
            # 空间分析图表 (13-18)
            "coordinate_matrix.png": {
                "title": "XYZ坐标对比矩阵（按目标力值分组）",
                "category": "空间分析",
                "description": "分析多变量间的两两关系和空间分布模式",
                "interpretation": """
• 对角线：各变量的密度分布
• 下三角：散点图
• 上三角：相关系数
• 颜色：目标力值分组
关注要点：
- 变量间的线性相关关系
- 异常点在多维空间的表现
- 相关系数的强度和方向
                """
            },
            "xy_heatmap.png": {
                "title": "XY平面密度热力图（按目标力值分组）",
                "category": "空间分析",
                "description": "显示数据点在平面上的密度分布和质量热点",
                "interpretation": """
• 等高线：数据密度等级
• 点的颜色和形状：合格性状态
• 密度高的区域：数据集中区
关注要点：
- 数据采集的空间代表性
- 高密度区域的质量表现
- 异常点的空间分布特征
                """
            },
            "parallel_coordinates.png": {
                "title": "并行坐标图 - 多维异常模式",
                "category": "空间分析",
                "description": "在多维空间中可视化数据模式和异常检测",
                "interpretation": """
• 各轴：不同的测量维度
• 连线：每个数据点的多维特征
• 颜色：目标力值分组
关注要点：
- 多维空间中的异常模式
- 不同维度间的协同变化
- 异常数据的多维特征
                """
            },
            "projection_combined.png": {
                "title": "2D投影图组合",
                "category": "空间分析",
                "description": "三维数据在不同平面上的投影分析",
                "interpretation": """
• XY投影：水平平面视图
• XZ投影：纵向视图
• YZ投影：侧向视图
• 点大小：绝对偏差大小
关注要点：
- 不同投影面上的异常分布
- 空间异常的方向性特征
- 三维数据的二维表现
                """
            },
            "spatial_clustering.png": {
                "title": "空间聚类异常检测图",
                "category": "空间分析",
                "description": "基于空间位置的聚类分析和异常检测",
                "interpretation": """
• 颜色：空间聚类分组
• 形状：实心圆（合格）vs 空心圆（异常）
• 大小：绝对偏差的大小
关注要点：
- 异常点在空间中的聚集模式
- 不同区域的质量表现差异
- 空间相关的质量问题
                """
            },
            "position_heatmap.png": {
                "title": "位置异常率热力图",
                "category": "空间分析",
                "description": "识别空间位置与质量表现的关系",
                "interpretation": """
• 颜色深浅：异常率高低
• 绿色：质量良好区域
• 黄色：质量一般区域
• 红色：质量问题区域
关注要点：
- 是否存在质量热点区域
- 空间异常模式的规律性
- 不同目标力值在相同位置的表现
                """
            },
            
            # 高级分析图表 (19-24)
            "correlation_matrix.png": {
                "title": "变量相关性矩阵",
                "category": "高级分析",
                "description": "显示各变量间的线性相关关系强度",
                "interpretation": """
• 颜色深浅：相关系数的强度
• 红色：正相关
• 蓝色：负相关
• 数值：具体的相关系数
关注要点：
- 哪些变量间存在强相关
- 相关关系的方向（正/负）
- 异常的相关模式
                """
            },
            "pareto_analysis.png": {
                "title": "帕雷托图 - 异常原因分析",
                "category": "高级分析",
                "description": "识别主要的异常原因，应用80/20原则",
                "interpretation": """
• 柱状图：各类异常的数量
• 折线图：累积百分比
• 异常类型：双重超差 > 绝对超差 > 百分比超差
关注要点：
- 哪种异常类型最常见
- 前80%的问题由哪几种原因造成
- 不同目标力值的异常模式差异
                """
            },
            "residual_analysis.png": {
                "title": "残差分析图",
                "category": "高级分析",
                "description": "检验模型假设和识别系统性误差",
                "interpretation": """
• X轴：模型拟合值（均值）
• Y轴：残差（实际值-拟合值）
• 红虚线：零残差线
• 平滑曲线：残差趋势
关注要点：
- 残差是否随机分布在零线两侧
- 是否存在残差的系统性模式
- 方差是否均匀（等方差性）
                """
            },
            "qq_plot.png": {
                "title": "QQ图 - 正态性检验",
                "category": "高级分析",
                "description": "检验数据是否符合正态分布假设",
                "interpretation": """
• X轴：理论正态分位数
• Y轴：样本分位数
• 直线：完美正态分布的参考线
关注要点：
- 数据点是否紧贴参考直线
- 尾部的偏离情况（重尾或轻尾）
- 整体分布的偏斜程度
                """
            },
            "radar_chart.png": {
                "title": "质量指标雷达图",
                "category": "高级分析",
                "description": "多指标综合评估的直观展示",
                "interpretation": """
• 各轴：不同的质量指标（成功率、Cp、Cpk）
• 距离中心的远近：指标得分高低
• 封闭图形的面积：综合质量水平
关注要点：
- 各指标的均衡发展情况
- 短板指标的识别
- 不同目标力值组的综合对比
                """
            },
            "waterfall_chart.png": {
                "title": "质量损失瀑布图",
                "category": "高级分析",
                "description": "累积质量损失的层次分析",
                "interpretation": """
• 柱状图：各目标力值的损失率
• 阶梯线：累积损失趋势
• 数值标签：具体的损失百分比
关注要点：
- 各目标力值的质量贡献
- 累积损失的构成
- 主要损失来源的识别
                """
            },
            
            # 过程能力与成功率分析 (25-27)
            "success_rate.png": {
                "title": "成功率趋势分析",
                "category": "过程能力",
                "description": "监控质量表现的时间趋势变化",
                "interpretation": """
• X轴：时间批次
• Y轴：成功率百分比
• 绿虚线：90%质量基准
• 蓝虚线：95%优秀基准
关注要点：
- 成功率的趋势方向
- 是否达到质量基准
- 批次间的稳定性
- 异常批次的识别
                """
            },
            "process_capability.png": {
                "title": "过程能力指数图",
                "category": "过程能力",
                "description": "评估过程满足规格要求的能力",
                "interpretation": """
• Cp：过程潜在能力（仅考虑变异）
• Cpk：过程实际能力（考虑偏移）
• 橙线：1.0（合格线）
• 绿线：1.33（优秀线）
关注要点：
- Cp ≥ 1.33: 过程能力优秀
- 1.0 ≤ Cp < 1.33: 过程能力合格
- Cpk显著小于Cp: 存在系统偏移
                """
            },
            "capability_histogram.png": {
                "title": "过程能力分析直方图",
                "category": "过程能力",
                "description": "直观显示过程分布与规格限制的关系",
                "interpretation": """
• 直方图：实际数据分布
• 曲线：正态密度拟合
• 红虚线：上下规格限制(USL/LSL)
• 绿实线：目标值
• 蓝点线：过程均值
关注要点：
- 分布是否完全在规格限制内
- 过程均值与目标值的偏离
- Cp≥1.33且Cpk≥1.33为优秀过程
                """
            },
            
            # 综合质量仪表盘 (28-30)
            "quality_dashboard.png": {
                "title": "质量控制仪表盘",
                "category": "质量仪表盘",
                "description": "关键质量指标的快速监控面板",
                "interpretation": """
• 成功率表盘：直观显示合格率
• 变异系数表盘：显示过程稳定性
• 数值标签：精确的指标值
关注要点：
- 成功率是否达到预期目标
- 变异系数是否在可接受范围
- 不同目标力值的表现对比
                """
            },
            "success_rate_trend.png": {
                "title": "成功率趋势详细分析",
                "category": "质量仪表盘",
                "description": "深入分析成功率的变化模式和预测",
                "interpretation": """
• 时间轴：详细的时间趋势
• 多条线：不同目标力值的成功率
• 置信区间：预测的不确定性
关注要点：
- 长期趋势的稳定性
- 短期波动的原因
- 未来趋势的预测
                """
            },
            
            # 新增高级分析图表 (31-36)
            "spatial_correlation_matrix.png": {
                "title": "误差与坐标相关性矩阵",
                "category": "空间分析",
                "description": "显示误差与各坐标轴的线性相关关系强度",
                "interpretation": """
• 红色：正相关（误差随坐标增大而增大）
• 蓝色：负相关（误差随坐标增大而减小）
• 数值：相关系数，越接近±1相关性越强
关注要点：
- 哪个坐标轴与误差关系最强
- 是否存在系统性的空间偏移
- 不同目标力值的空间规律是否一致
                """
            },
            "error_spatial_distribution.png": {
                "title": "误差空间分布图（XY平面）",
                "category": "空间分析",
                "description": "在XY平面上显示误差的空间分布模式",
                "interpretation": """
• 颜色：绿色=误差小，红色=误差大
• 大小：点的大小表示误差值大小
• 分面：不同目标力值的独立分析
关注要点：
- 是否存在误差聚集的热点区域
- 空间分布是否均匀随机
- 不同目标力值的空间表现差异
                """
            },
            "error_distribution_analysis.png": {
                "title": "误差分布特性分析",
                "category": "误差分布分析",
                "description": "分析误差分布是否符合正态分布假设",
                "interpretation": """
• 直方图：误差的实际频次分布
• 蓝线：实际密度曲线
• 红虚线：理论正态分布拟合
关注要点：
- 实际分布与正态分布的拟合程度
- 是否存在偏斜或多峰分布
- 分布形状是否暗示特殊原因变异
                """
            },
            "error_qq_plot.png": {
                "title": "误差分布QQ图",
                "category": "误差分布分析",
                "description": "检验误差是否符合正态分布",
                "interpretation": """
• X轴：理论正态分位数
• Y轴：样本分位数
• 红线：完美正态分布的参考线
关注要点：
- 数据点是否紧贴参考直线
- 尾部偏离表示重尾或轻尾分布
- 整体偏离表示分布偏斜
                """
            },
            "machine_performance_comparison.png": {
                "title": "各机台性能对比（成功率）",
                "category": "多源变异分析",
                "description": "对比不同机台的质量表现，识别设备相关问题",
                "interpretation": """
• 柱状图：各机台的成功率
• 橙线：90%质量基准
• 绿线：95%优秀基准
• 数值标签：精确的成功率
关注要点：
- 哪台机台表现最好/最差
- 机台间差异是否显著
- 是否需要针对性设备维护
                """
            },
            "shift_performance_comparison.png": {
                "title": "各班次误差对比（平均误差）",
                "category": "多源变异分析", 
                "description": "对比不同班次的精度表现，识别人员或时间相关问题",
                "interpretation": """
• 柱状图：各班次的平均误差
• 数值越小表示精度越高
• 分面：不同目标力值的独立分析
关注要点：
- 班次间误差是否存在系统性差异
- 是否存在特定班次的问题
- 人员培训或设备调试需求
                """
            }
        }
        
        # 收集所有存在的图表文件
        chart_files = [f for f in output_dir.glob('*.png') if f.is_file()]
        
        for chart_path in chart_files:
            chart_name = chart_path.name
            if chart_name in chart_mapping:
                chart_info = chart_mapping[chart_name]
                chart_data = {
                    'chart_id': chart_name.split('.')[0],
                    'title': chart_info['title'],
                    'category': chart_info['category'],
                    'description': chart_info['description'],
                    'interpretation': chart_info['interpretation'],
                    'file_path': str(chart_path),
                    'file_url': f"/static/charts/{task_id}/{chart_name}",
                    'chart_type': 'analysis',
                    'filename': chart_name
                }
                charts.append(chart_data)
        
        # 按类别排序
        category_order = ["基础分析", "偏差分析", "统计过程控制", "空间分析", "高级分析", "过程能力", "质量仪表盘", "误差分布分析", "多源变异分析"]
        charts.sort(key=lambda x: (category_order.index(x['category']) if x['category'] in category_order else 999, x['title']))
        
        return charts
    
    def _collect_files(self, task_id: str, output_dir: Path) -> Dict[str, Any]:
        """收集生成的数据文件"""
        files = {}
        
        # 清理后的数据
        cleaned_data_file = output_dir / "cleaned_data.csv"
        if cleaned_data_file.exists():
            files['cleaned_data'] = {
                'filename': 'cleaned_data.csv',
                'path': str(cleaned_data_file),
                'url': f"/static/charts/{task_id}/cleaned_data.csv",
                'description': '清理后的原始数据'
            }
        
        # 3D散点图数据
        scatter_data_file = output_dir / "scatter_3d_data.csv"
        if scatter_data_file.exists():
            files['scatter_3d_data'] = {
                'filename': 'scatter_3d_data.csv',
                'path': str(scatter_data_file),
                'url': f"/static/charts/{task_id}/scatter_3d_data.csv",
                'description': '3D散点图数据'
            }
        
        # 分析结果JSON
        results_file = output_dir / "analysis_results.json"
        if results_file.exists():
            files['analysis_results'] = {
                'filename': 'analysis_results.json',
                'path': str(results_file),
                'url': f"/static/charts/{task_id}/analysis_results.json",
                'description': '详细分析结果'
            }
        
        return files
    
    def _generate_summary_report(self, statistics: Dict) -> str:
        """生成分析摘要报告"""
        try:
            summary = statistics.get('summary', {})
            overall = statistics.get('overall_stats', [])
            
            # 处理overall_stats可能是数组的情况
            if isinstance(overall, list) and len(overall) > 0:
                overall = overall[0]
            elif not isinstance(overall, dict):
                overall = {}
            
            total_records = summary.get('total_records', overall.get('样本数', overall.get('total_records', 0)))
            success_rate = summary.get('success_rate', overall.get('success_rate', 0))
            mean_force = summary.get('mean_force', overall.get('均值', overall.get('mean_force', 0)))
            
            # 获取过程能力信息
            process_capability = statistics.get('process_capability', [])
            capability_summary = ""
            if process_capability and isinstance(process_capability, list) and len(process_capability) > 0:
                capability_summary = "\n\n过程能力概览:\n"
                for capability in process_capability:
                    if isinstance(capability, dict):
                        目标力值 = capability.get('target_force', 0)
                        Cp = capability.get('Cp', 0)
                        Cpk = capability.get('Cpk', 0)
                        等级 = capability.get('能力等级', '未知')
                        capability_summary += f"- 目标{目标力值}N: Cp={Cp:.3f}, Cpk={Cpk:.3f}, 等级={等级}\n"
            else:
                logger.warning(f"摘要报告：过程能力数据为空: {process_capability}")
            
            report = f"""
数据分析摘要报告

数据概览:
- 总数据点: {total_records}
- 平均力值: {mean_force:.2f} N
- 整体成功率: {success_rate:.1f}%{capability_summary}

分析完成时间: {datetime.now(SHANGHAI_TZ).strftime('%Y-%m-%d %H:%M:%S %Z')}
            """.strip()
            
            return report
            
        except Exception as e:
            logger.error(f"生成摘要报告失败: {str(e)}")
            return f"报告生成失败: {str(e)}"
    
    def _generate_quality_report(self, statistics: Dict) -> str:
        """生成数据质量报告"""
        try:
            overall = statistics.get('overall_stats', [])
            summary = statistics.get('summary', {})
            
            # 处理overall_stats可能是数组的情况
            if isinstance(overall, list) and len(overall) > 0:
                overall = overall[0]
            elif not isinstance(overall, dict):
                overall = {}
                
            success_rate = summary.get('success_rate', overall.get('success_rate', 0))
            
            if success_rate >= 95:
                quality_level = "优秀"
            elif success_rate >= 90:
                quality_level = "良好"
            elif success_rate >= 80:
                quality_level = "一般"
            else:
                quality_level = "需要改进"
            
            return f"数据质量等级: {quality_level} (成功率: {success_rate:.1f}%)"
            
        except Exception as e:
            return f"质量报告生成失败: {str(e)}"
    
    def _generate_recommendations(self, statistics: Dict) -> List[str]:
        """生成改进建议"""
        try:
            recommendations = []
            overall = statistics.get('overall_stats', [])
            summary = statistics.get('summary', {})
            
            # 处理overall_stats可能是数组的情况
            if isinstance(overall, list) and len(overall) > 0:
                overall = overall[0]
            elif not isinstance(overall, dict):
                overall = {}
                
            success_rate = summary.get('success_rate', overall.get('success_rate', 0))
            
            if success_rate < 90:
                recommendations.append("建议检查测量系统的稳定性")
                recommendations.append("考虑调整工艺参数以提高精度")
            
            if success_rate < 80:
                recommendations.append("建议进行设备校准")
                recommendations.append("检查环境条件是否影响测量")
            
            # 检查变异系数
            cv_percent = summary.get('cv_percent', overall.get('变异系数', 0))
            if cv_percent > 10:
                recommendations.append("数据变异性较大，建议优化工艺稳定性")
            
            # 检查异常值情况
            outlier_summary = statistics.get('outlier_summary', [])
            if outlier_summary and isinstance(outlier_summary, list):
                high_outlier_rate = False
                for outlier in outlier_summary:
                    if isinstance(outlier, dict):
                        iqr_rate = outlier.get('IQR异常率', 0)
                        z_rate = outlier.get('Z异常率', 0)
                        if iqr_rate > 5 or z_rate > 2:
                            high_outlier_rate = True
                            break
                
                if high_outlier_rate:
                    recommendations.append("检测到较多异常值，建议检查数据采集过程")
            
            if len(recommendations) == 0:
                recommendations.append("系统运行良好，继续保持")
            
            return recommendations
            
        except Exception as e:
            logger.error(f"生成改进建议失败: {str(e)}")
            return [f"建议生成失败: {str(e)}"]

    def _generate_analysis_interpretation(self, statistics: Dict) -> Dict[str, str]:
        """生成详细分析解读"""
        try:
            interpretation = {}
            
            # 1. 数据质量解读
            data_summary = statistics.get('data_summary', [])
            if data_summary and isinstance(data_summary, list) and len(data_summary) > 0:
                # 从数组中取第一个元素（R生成的JSON格式）
                data_summary_dict = data_summary[0] if isinstance(data_summary[0], dict) else {}
                quality_text = f"""数据质量检查报告：
• 总数据行数：{data_summary_dict.get('总行数', 0)}
• 缺失值数量：{data_summary_dict.get('缺失值', 0)}
• 重复行数量：{data_summary_dict.get('重复行', 0)}
• 力值范围：{data_summary_dict.get('力值最小值', 0):.2f} - {data_summary_dict.get('力值最大值', 0):.2f} N
• 平均力值：{data_summary_dict.get('力值均值', 0):.2f} N (标准差: {data_summary_dict.get('力值标准差', 0):.2f} N)

数据质量评估："""
                if data_summary_dict.get('缺失值', 0) == 0 and data_summary_dict.get('重复行', 0) == 0:
                    quality_text += "数据完整性良好，无缺失值和重复数据。"
                else:
                    quality_text += "发现数据质量问题，建议进行数据清理。"
                
                interpretation['data_quality'] = quality_text
            elif isinstance(data_summary, dict):
                # 如果直接是字典格式
                quality_text = f"""数据质量检查报告：
• 总数据行数：{data_summary.get('总行数', 0)}
• 缺失值数量：{data_summary.get('缺失值', 0)}
• 重复行数量：{data_summary.get('重复行', 0)}
• 力值范围：{data_summary.get('力值最小值', 0):.2f} - {data_summary.get('力值最大值', 0):.2f} N
• 平均力值：{data_summary.get('力值均值', 0):.2f} N (标准差: {data_summary.get('力值标准差', 0):.2f} N)

数据质量评估："""
                if data_summary.get('缺失值', 0) == 0 and data_summary.get('重复行', 0) == 0:
                    quality_text += "数据完整性良好，无缺失值和重复数据。"
                else:
                    quality_text += "发现数据质量问题，建议进行数据清理。"
                
                interpretation['data_quality'] = quality_text
            
            # 2. 趋势分析解读
            trend_stats = statistics.get('trend_stats', [])
            if trend_stats and isinstance(trend_stats, list) and len(trend_stats) > 0:
                trend_text = "趋势分析结果：\n"
                for trend in trend_stats:
                    if isinstance(trend, dict):
                        分组 = trend.get('分组', trend.get('target_force', ''))
                        斜率 = trend.get('斜率', trend.get('slope', 0))
                        p值 = trend.get('p值', trend.get('p_value', 1))
                        R平方 = trend.get('R平方', trend.get('r_squared', 0))
                        
                        if p值 < 0.05:
                            方向 = "上升" if 斜率 > 0 else "下降"
                            trend_text += f"• {分组}：检测到显著{方向}趋势（斜率={斜率:.6f}, p值={p值:.4f}, R²={R平方:.4f}）\n"
                            trend_text += f"  - 每增加一个序号，力值平均变化{abs(斜率):.6f}N\n"
                        else:
                            trend_text += f"• {分组}：无显著趋势（p值={p值:.4f}）\n"
                
                interpretation['trend_analysis'] = trend_text
            
            # 3. 异常值分析解读
            outlier_summary = statistics.get('outlier_summary', [])
            if outlier_summary and isinstance(outlier_summary, list) and len(outlier_summary) > 0:
                outlier_text = "异常值检测结果：\n"
                for outlier in outlier_summary:
                    if isinstance(outlier, dict):
                        目标力值 = outlier.get('target_force', 0)
                        IQR异常率 = outlier.get('IQR异常率', outlier.get('iqr_outlier_rate', 0))
                        Z异常率 = outlier.get('Z异常率', outlier.get('z_outlier_rate', 0))
                        
                        outlier_text += f"• 目标{目标力值}N组：\n"
                        outlier_text += f"  - IQR方法检测异常率：{IQR异常率}%\n"
                        outlier_text += f"  - Z-score方法检测异常率：{Z异常率}%\n"
                        
                        if IQR异常率 > 5 or Z异常率 > 1:
                            outlier_text += "  - 建议：异常值比例偏高，需要检查测量系统稳定性\n"
                
                interpretation['outlier_analysis'] = outlier_text
            
            # 4. 稳定性分析解读
            stability_analysis = statistics.get('stability_analysis', [])
            if stability_analysis and isinstance(stability_analysis, list) and len(stability_analysis) > 0:
                stability_text = "稳定性分析（游程检验）：\n"
                for stability in stability_analysis:
                    if isinstance(stability, dict):
                        目标力值 = stability.get('target_force', 0)
                        总游程数 = stability.get('总游程数', stability.get('total_runs', 0))
                        平均游程长度 = stability.get('平均游程长度', stability.get('avg_run_length', 0))
                        最长游程 = stability.get('最长游程', stability.get('max_run_length', 0))
                        
                        stability_text += f"• 目标{目标力值}N组：\n"
                        stability_text += f"  - 总游程数：{总游程数}（游程越多表示数据变化越频繁）\n"
                        stability_text += f"  - 平均游程长度：{平均游程长度}\n"
                        stability_text += f"  - 最长游程：{最长游程}\n"
                        
                        if 最长游程 > 10:
                            stability_text += "  - 注意：存在较长的连续游程，可能存在系统性偏移\n"
                
                interpretation['stability_analysis'] = stability_text
            
            # 5. 变化点检测解读
            change_points = statistics.get('change_point_analysis', [])
            if change_points and isinstance(change_points, list) and len(change_points) > 0:
                change_text = "变化点检测结果：\n"
                has_change = False
                for change in change_points:
                    if isinstance(change, dict):
                        目标力值 = change.get('target_force', 0)
                        变化点数量 = change.get('潜在变化点数量', change.get('change_points', 0))
                        最大变化 = change.get('最大均值变化', change.get('max_change', 0))
                        
                        if 变化点数量 > 0:
                            has_change = True
                            change_text += f"• 目标{目标力值}N组：检测到{变化点数量}个潜在变化点\n"
                            change_text += f"  - 最大均值变化：{最大变化:.3f}N\n"
                            change_text += "  - 建议：检查这些时间点是否有工艺调整或设备变化\n"
                
                if not has_change:
                    change_text += "• 未检测到明显的过程变化点，过程相对稳定\n"
                
                interpretation['change_point'] = change_text
            
            # 6. 自相关分析解读
            autocorr_analysis = statistics.get('autocorr_analysis', [])
            if autocorr_analysis and isinstance(autocorr_analysis, list) and len(autocorr_analysis) > 0:
                autocorr_text = "自相关分析结果：\n"
                for autocorr in autocorr_analysis:
                    if isinstance(autocorr, dict):
                        目标力值 = autocorr.get('target_force', 0)
                        lag1 = autocorr.get('lag1_correlation', 0)
                        lag2 = autocorr.get('lag2_correlation', 0)
                        lag3 = autocorr.get('lag3_correlation', 0)
                        
                        autocorr_text += f"• 目标{目标力值}N组：\n"
                        autocorr_text += f"  - 滞后1期相关性：{lag1:.4f}\n"
                        autocorr_text += f"  - 滞后2期相关性：{lag2:.4f}\n"
                        autocorr_text += f"  - 滞后3期相关性：{lag3:.4f}\n"
                        
                        if abs(lag1) > 0.7:
                            autocorr_text += "  - 注意：存在强自相关，相邻测量值高度相关\n"
                
                interpretation['autocorrelation'] = autocorr_text
            
            # 7. 过程能力分析解读
            process_capability = statistics.get('process_capability', [])
            logger.info(f"过程能力数据: {process_capability}")
            
            if process_capability and isinstance(process_capability, list) and len(process_capability) > 0:
                capability_text = "过程能力分析：\n"
                for capability in process_capability:
                    if isinstance(capability, dict):
                        目标力值 = capability.get('target_force', 0)
                        Cp = capability.get('Cp', 0)
                        Cpk = capability.get('Cpk', 0)
                        等级 = capability.get('能力等级', '未知')  # 直接使用R脚本生成的字段
                        
                        capability_text += f"• 目标{目标力值}N组：\n"
                        capability_text += f"  - Cp = {Cp:.3f}（过程潜在能力）\n"
                        capability_text += f"  - Cpk = {Cpk:.3f}（过程实际能力）\n"
                        capability_text += f"  - 能力等级：{等级}\n"
                        
                        # 解释Cp和Cpk的含义
                        if Cp > 0 and Cpk > 0 and Cpk < Cp * 0.8:
                            capability_text += "  - 注意：Cpk明显小于Cp，说明过程存在偏移\n"
                        
                        # 根据等级给出建议
                        if 等级 == "不合格":
                            capability_text += "  - 建议：过程能力不足，需要大幅改进\n"
                        elif 等级 == "勉强":
                            capability_text += "  - 建议：过程能力有限，建议优化\n"
                        elif 等级 == "合格":
                            capability_text += "  - 建议：过程能力良好，可以继续优化\n"
                        elif 等级 == "优秀":
                            capability_text += "  - 评价：过程能力优秀，保持现状\n"
                
                # 添加数据来源说明，确保一致性
                capability_text += "\n注：以上数据与过程能力图表中显示的数值保持一致。"
                interpretation['process_capability'] = capability_text
            else:
                # 添加调试信息
                logger.warning(f"过程能力数据为空或格式错误: {process_capability}")
                interpretation['process_capability'] = "过程能力分析数据不可用"
            
            return interpretation
            
        except Exception as e:
            logger.error(f"生成分析解读失败: {str(e)}")
            import traceback
            logger.error(f"详细错误: {traceback.format_exc()}")
            return {
                'error': f"分析解读生成失败: {str(e)}"
            } 

    def get_analysis_history(self) -> List[Dict[str, Any]]:
        """获取所有分析历史记录"""
        history_path = Path(settings.HISTORY_DIR)
        history_path.mkdir(exist_ok=True)
        history_files = sorted(history_path.glob("*.json"), key=os.path.getmtime, reverse=True)

        history = []
        for file in history_files:
            try:
                task_id = file.stem
                record = json.loads(file.read_text(encoding='utf-8'))

                # 尝试从关联的 analysis_results.json 读取更精确的统计数据
                try:
                    results_file = Path(settings.CHARTS_DIR) / task_id / "analysis_results.json"
                    if results_file.exists():
                        results_data = json.loads(results_file.read_text(encoding='utf-8'))
                        target_analysis = results_data.get('target_analysis', [])
                        
                        if target_analysis and isinstance(target_analysis, list):
                            total_points = sum(item.get('数据点数', 0) for item in target_analysis if item.get('数据点数') is not None)
                            if total_points > 0:
                                weighted_success_sum = sum(
                                    item.get('成功率_综合', 0) * item.get('数据点数', 0) 
                                    for item in target_analysis 
                                    if item.get('成功率_综合') is not None and item.get('数据点数') is not None
                                )
                                record['successRate'] = round(weighted_success_sum / total_points, 2)
                except Exception as e:
                    logger.warning(f"为任务 {task_id} 计算精确成功率失败: {e}。将使用历史文件中的值。")

                history.append(record)
            except Exception as e:
                logger.warning(f"无法读取或处理历史文件 {file.name}: {e}")
                continue
        
        return history

    def clear_all_history(self):
        """清空所有历史记录和相关的分析文件"""
        history_path = Path(settings.HISTORY_DIR)
        charts_path = Path(settings.CHARTS_DIR)
        
        deleted_count = 0
        
        # 删除历史记录json文件
        if history_path.exists():
            for f in history_path.glob("*.json"):
                try:
                    f.unlink()
                    deleted_count += 1
                except Exception as e:
                    logger.error(f"删除历史文件失败 {f}: {e}")

        # 删除分析结果目录
        if charts_path.exists():
            for task_dir in charts_path.iterdir():
                if task_dir.is_dir():
                    try:
                        shutil.rmtree(task_dir)
                    except Exception as e:
                        logger.error(f"删除分析目录失败 {task_dir}: {e}")
        
        logger.info(f"成功清空历史记录。删除了 {deleted_count} 条记录和相关的分析目录。")
        return deleted_count

    def generate_word_report(self, task_id: str) -> Optional[str]:
        """根据分析结果生成完整的Word报告，包含所有图表和详细解读。如果存在DeepSeek分析，也会一并包含"""
        # 结果文件路径
        results_file = Path(settings.CHARTS_DIR) / task_id / "analysis_results.json"
        
        if not results_file.exists():
            logger.error(f"报告生成失败: 未找到任务 {task_id} 的结果文件")
            return None

        with open(results_file, 'r', encoding='utf-8') as f:
            data = json.load(f)

        # 检查是否存在DeepSeek分析结果
        deepseek_file = Path(settings.CHARTS_DIR) / task_id / "deepseek_analysis.json"
        deepseek_content = None
        
        if deepseek_file.exists():
            try:
                with open(deepseek_file, 'r', encoding='utf-8') as f:
                    deepseek_data = json.load(f)
                    deepseek_content = deepseek_data.get('analysis_report', '')
                logger.info(f"找到DeepSeek分析结果，将包含在Word报告中")
            except Exception as e:
                logger.warning(f"读取DeepSeek分析文件失败: {e}")

        # 创建Word文档
        doc = Document()
        
        # 设置文档样式
        style = doc.styles['Normal']
        style.font.name = 'SimSun'
        style.font.size = Pt(11)
        
        # 添加标题
        if deepseek_content:
            title = doc.add_heading('压力系统数据综合分析报告', 0)
            subtitle = doc.add_paragraph()
            subtitle_run = subtitle.add_run('包含DeepSeek AI智能分析与R统计分析')
            subtitle_run.font.size = Pt(14)
            subtitle_run.font.italic = True
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            title = doc.add_heading('压力系统数据分析综合报告', 0)
        
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加生成时间和基本信息
        doc.add_paragraph(f'报告生成时间：{datetime.now(SHANGHAI_TZ).strftime("%Y年%m月%d日 %H:%M:%S %Z")}')
        doc.add_paragraph(f'任务ID：{task_id}')
        doc.add_paragraph('')

        # 如果存在DeepSeek分析，添加AI分析部分
        if deepseek_content:
            doc.add_heading('第一部分：DeepSeek AI智能分析', level=1)
            doc.add_paragraph('本部分基于人工智能技术，对压力系统数据进行深度分析和专业解读。')
            doc.add_paragraph('')
            
            # 将DeepSeek报告按段落分割并添加到Word文档
            deepseek_paragraphs = deepseek_content.split('\n')
            current_paragraph = ""
            
            for para in deepseek_paragraphs:
                para = para.strip()
                if not para:
                    if current_paragraph:
                        doc.add_paragraph(current_paragraph)
                        current_paragraph = ""
                    continue
                
                # 检查是否是标题（以#开头）
                if para.startswith('####'):
                    if current_paragraph:
                        doc.add_paragraph(current_paragraph)
                        current_paragraph = ""
                    doc.add_heading(para.replace('####', '').strip(), level=4)
                elif para.startswith('###'):
                    if current_paragraph:
                        doc.add_paragraph(current_paragraph)
                        current_paragraph = ""
                    doc.add_heading(para.replace('###', '').strip(), level=3)
                elif para.startswith('##'):
                    if current_paragraph:
                        doc.add_paragraph(current_paragraph)
                        current_paragraph = ""
                    doc.add_heading(para.replace('##', '').strip(), level=2)
                elif para.startswith('#'):
                    if current_paragraph:
                        doc.add_paragraph(current_paragraph)
                        current_paragraph = ""
                    doc.add_heading(para.replace('#', '').strip(), level=1)
                else:
                    # 普通段落
                    if current_paragraph:
                        current_paragraph += "\n" + para
                    else:
                        current_paragraph = para
            
            # 添加最后一个段落
            if current_paragraph:
                doc.add_paragraph(current_paragraph)
            
            doc.add_page_break()
            doc.add_heading('第二部分：R统计分析详细数据', level=1)
        else:
            doc.add_heading('第一部分：统计分析概览', level=1)
        
        # 1. 执行摘要
        doc.add_heading('1. 执行摘要', level=1)
        summary_text = self._generate_summary_report(data)
        doc.add_paragraph(summary_text)
        
        # 2. 数据质量评估
        doc.add_heading('2. 数据质量评估', level=1)
        interpretations = self._generate_analysis_interpretation(data)
        
        if 'data_quality' in interpretations:
            doc.add_paragraph(interpretations['data_quality'])
        
        # 3. 目标力值分析汇总
        doc.add_heading('3. 目标力值分析汇总', level=1)
        target_analysis = data.get('target_analysis', [])
        if target_analysis:
            target_table = doc.add_table(rows=1, cols=8)
            target_table.style = 'Table Grid'
            
            headers = ['目标力值(N)', '数据点数', '成功率(%)', '平均力值(N)', 
                      '平均偏差(N)', '标准差(N)', '最大偏差(N)', '容差限制(N)']
            hdr_cells = target_table.rows[0].cells
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                hdr_cells[i].paragraphs[0].runs[0].font.bold = True
            
            for target in target_analysis:
                if isinstance(target, dict):
                    row_cells = target_table.add_row().cells
                    row_cells[0].text = str(target.get('target_force', ''))
                    row_cells[1].text = str(target.get('数据点数', ''))
                    row_cells[2].text = f"{target.get('成功率_综合', 0):.1f}"
                    row_cells[3].text = f"{target.get('平均力值', 0):.2f}"
                    row_cells[4].text = f"{target.get('平均偏差_绝对', 0):.2f}"
                    row_cells[5].text = f"{target.get('标准差', 0):.2f}"
                    row_cells[6].text = f"{target.get('最大偏差_绝对', 0):.2f}"
                    row_cells[7].text = f"{target.get('绝对容差限制', 0):.2f}"

        # 4. 过程能力评估
        doc.add_heading('4. 过程能力评估', level=1)
        process_capability = data.get('process_capability', [])
        if process_capability:
            capability_table = doc.add_table(rows=1, cols=4)
            capability_table.style = 'Table Grid'
            
            headers = ['目标力值(N)', 'Cp', 'Cpk', '能力等级']
            hdr_cells = capability_table.rows[0].cells
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                hdr_cells[i].paragraphs[0].runs[0].font.bold = True
            
            for cap in process_capability:
                if isinstance(cap, dict):
                    row_cells = capability_table.add_row().cells
                    row_cells[0].text = str(cap.get('target_force', ''))
                    row_cells[1].text = f"{cap.get('Cp', 0):.3f}"
                    row_cells[2].text = f"{cap.get('Cpk', 0):.3f}"
                    row_cells[3].text = cap.get('能力等级', '')

        if 'process_capability' in interpretations:
            doc.add_paragraph('\n过程能力分析解读：')
            doc.add_paragraph(interpretations['process_capability'])

        # 5. 图表展示
        doc.add_heading('5. 详细图表分析', level=1)
        charts = self._collect_charts(task_id, Path(settings.CHARTS_DIR) / task_id)
        
        if charts:
            # 按类别组织图表
            categories = {}
            for chart in charts:
                category = chart['category']
                if category not in categories:
                    categories[category] = []
                categories[category].append(chart)
            
            # 显示所有图表，按类别分组
            category_order = ["基础分析", "偏差分析", "统计过程控制", "空间分析", "高级分析", "过程能力", "质量仪表盘"]
            
            chart_count = 0
            for category in category_order:
                if category in categories and categories[category]:
                    doc.add_heading(f'5.{list(category_order).index(category)+1} {category}', level=2)
                    
                    for chart in categories[category]:
                        chart_count += 1
                        
                        # 图表标题
                        chart_title = doc.add_paragraph()
                        chart_title_run = chart_title.add_run(f"图表 {chart_count}: {chart['title']}")
                        chart_title_run.font.size = Pt(12)
                        chart_title_run.font.bold = True
                        
                        # 插入图表
                        if os.path.exists(chart['file_path']):
                            try:
                                doc.add_picture(chart['file_path'], width=Inches(6.0))
                                
                                # 图表描述和解读
                                desc_para = doc.add_paragraph()
                                desc_run = desc_para.add_run(f"描述：{chart['description']}")
                                desc_run.font.size = Pt(10)
                                desc_run.font.italic = True
                                
                                if chart.get('interpretation'):
                                    interp_para = doc.add_paragraph()
                                    interp_run = interp_para.add_run("解读说明：")
                                    interp_run.font.size = Pt(10)
                                    interp_run.font.bold = True
                                    
                                    interp_content = doc.add_paragraph()
                                    interp_content_run = interp_content.add_run(chart['interpretation'])
                                    interp_content_run.font.size = Pt(9)
                                    
                            except Exception as e:
                                logger.error(f"插入图表失败: {chart['title']} - {e}")
                                doc.add_paragraph(f"❌ 图表插入失败: {chart['title']}")
                            else:
                                doc.add_paragraph(f"⚠️ 图表文件未找到: {chart['file_path']}")
                            
                            doc.add_paragraph('')

        # 6. 改进建议
        doc.add_heading('6. 改进建议', level=1)
        recommendations = self._generate_recommendations(data)
        for i, rec in enumerate(recommendations, 1):
            doc.add_paragraph(f"{i}. {rec}")

        # 7. 结论
        doc.add_heading('7. 分析结论', level=1)
        total_points = len(data.get('data_with_target', []))
        
        target_success_rates = []
        if target_analysis:
            for target in target_analysis:
                if isinstance(target, dict):
                    target_success_rates.append(target.get('成功率_综合', 0))
        
        success_rate = sum(target_success_rates) / len(target_success_rates) if target_success_rates else 0
        
        conclusion = f"""
本次分析共处理 {total_points} 个数据点，综合成功率为 {success_rate:.1f}%。

基于以上分析结果：
• 数据质量：{'优秀' if success_rate >= 95 else '良好' if success_rate >= 90 else '一般' if success_rate >= 80 else '需改进'}
• 图表总数：{len(charts)} 个，涵盖基础分析、控制图、空间分析等多个维度
• 分析深度：包含统计过程控制、异常检测、空间分析和过程能力评估
{'• AI智能分析：包含DeepSeek人工智能深度解读和专业建议' if deepseek_content else ''}

建议定期执行此类分析以持续监控系统性能和质量水平。
        """
        doc.add_paragraph(conclusion)
        
        # 添加页脚信息
        doc.add_paragraph()
        footer_para = doc.add_paragraph()
        footer_run = footer_para.add_run(f"报告由压力分析系统自动生成 | 任务ID: {task_id} | 生成时间: {datetime.now(SHANGHAI_TZ).strftime('%Y-%m-%d %H:%M:%S %Z')}")
        footer_run.font.size = Pt(9)
        footer_run.font.italic = True
        
        # 保存报告到两个位置
        # 1. temp/reports 目录 (临时下载)
        temp_reports_dir = Path(__file__).parent.parent.parent / "temp" / "reports"
        temp_reports_dir.mkdir(parents=True, exist_ok=True)
        temp_report_path = temp_reports_dir / f"pressure_analysis_comprehensive_report_{task_id}.docx"
        
        # 2. output/reports 目录 (永久存储)
        output_reports_dir = Path(settings.CHARTS_DIR).parent / "reports"
        output_reports_dir.mkdir(parents=True, exist_ok=True)
        output_report_path = output_reports_dir / f"pressure_analysis_comprehensive_report_{task_id}.docx"
        
        try:
            # 保存到临时目录
            doc.save(temp_report_path)
            logger.info(f"综合报告已保存到临时目录: {temp_report_path}")
            
            # 复制到输出目录
            import shutil
            shutil.copy2(temp_report_path, output_report_path)
            logger.info(f"综合报告已保存到输出目录: {output_report_path}")
            
            return str(temp_report_path)
        except Exception as e:
            logger.error(f"保存Word报告失败: {e}")
            return None

    def generate_comprehensive_word_report(self, task_id: str, analysis_data: Dict[str, Any], deepseek_report: str) -> Optional[str]:
        """生成包含DeepSeek分析和R分析结果的综合Word报告"""
        try:
            # 创建Word文档
            doc = Document()
            
            # 设置文档样式
            style = doc.styles['Normal']
            style.font.name = 'SimSun'
            style.font.size = Pt(11)
            
            # 添加标题
            title = doc.add_heading('综合压力系统数据分析报告', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加副标题
            subtitle = doc.add_paragraph()
            subtitle_run = subtitle.add_run('基于DeepSeek AI分析与R统计分析的综合评估报告')
            subtitle_run.font.size = Pt(14)
            subtitle_run.font.italic = True
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加生成时间和基本信息
            doc.add_paragraph('')
            doc.add_paragraph(f'报告生成时间：{datetime.now(SHANGHAI_TZ).strftime("%Y年%m月%d日 %H:%M:%S %Z")}')
            doc.add_paragraph(f'任务ID：{task_id}')
            doc.add_paragraph(f'分析方法：DeepSeek AI智能分析 + R统计分析')
            doc.add_paragraph('')
            
            # ============================================================================
            # 第一部分：DeepSeek AI智能分析报告
            # ============================================================================
            doc.add_heading('第一部分：DeepSeek AI智能分析报告', level=1)
            
            # 将DeepSeek报告按段落分割并添加到Word文档
            deepseek_paragraphs = deepseek_report.split('\n')
            current_paragraph = ""
            
            for para in deepseek_paragraphs:
                para = para.strip()
                if not para:
                    if current_paragraph:
                        doc.add_paragraph(current_paragraph)
                        current_paragraph = ""
                    continue
                
                # 检查是否是标题（以#开头）
                if para.startswith('####'):
                    if current_paragraph:
                        doc.add_paragraph(current_paragraph)
                        current_paragraph = ""
                    doc.add_heading(para.replace('####', '').strip(), level=4)
                elif para.startswith('###'):
                    if current_paragraph:
                        doc.add_paragraph(current_paragraph)
                        current_paragraph = ""
                    doc.add_heading(para.replace('###', '').strip(), level=3)
                elif para.startswith('##'):
                    if current_paragraph:
                        doc.add_paragraph(current_paragraph)
                        current_paragraph = ""
                    doc.add_heading(para.replace('##', '').strip(), level=2)
                elif para.startswith('#'):
                    if current_paragraph:
                        doc.add_paragraph(current_paragraph)
                        current_paragraph = ""
                    doc.add_heading(para.replace('#', '').strip(), level=1)
                else:
                    # 普通段落
                    if current_paragraph:
                        current_paragraph += "\n" + para
                    else:
                        current_paragraph = para
            
            # 添加最后一个段落
            if current_paragraph:
                doc.add_paragraph(current_paragraph)
            
            # ============================================================================
            # 第二部分：R统计分析详细数据
            # ============================================================================
            doc.add_page_break()
            doc.add_heading('第二部分：R统计分析详细数据与图表', level=1)
            
            # 2.1 数据概览
            doc.add_heading('2.1 数据概览', level=2)
            data_summary = analysis_data.get('data_summary', [])
            
            # 处理 data_summary 可能是列表的情况
            if data_summary:
                if isinstance(data_summary, list) and len(data_summary) > 0:
                    summary_data = data_summary[0]  # 取第一个元素
                elif isinstance(data_summary, dict):
                    summary_data = data_summary
                else:
                    summary_data = {}
                
                if summary_data:
                    summary_table = doc.add_table(rows=1, cols=2)
                    summary_table.style = 'Table Grid'
                    
                    hdr_cells = summary_table.rows[0].cells
                    hdr_cells[0].text = '指标'
                    hdr_cells[1].text = '数值'
                    
                    for key, value in summary_data.items():
                        row_cells = summary_table.add_row().cells
                        row_cells[0].text = str(key)
                        row_cells[1].text = str(value)
            
            # 2.2 整体统计指标
            doc.add_heading('2.2 整体统计指标', level=2)
            overall_stats = analysis_data.get('overall_stats', [])
            
            # 处理 overall_stats 可能是列表的情况
            if overall_stats:
                if isinstance(overall_stats, list) and len(overall_stats) > 0:
                    stats_data = overall_stats[0]  # 取第一个元素
                elif isinstance(overall_stats, dict):
                    stats_data = overall_stats
                else:
                    stats_data = {}
                
                if stats_data:
                    stats_table = doc.add_table(rows=1, cols=2)
                    stats_table.style = 'Table Grid'
                    
                    hdr_cells = stats_table.rows[0].cells
                    hdr_cells[0].text = '统计指标'
                    hdr_cells[1].text = '数值'
                    
                    for key, value in stats_data.items():
                        row_cells = stats_table.add_row().cells
                        row_cells[0].text = str(key)
                        if isinstance(value, (int, float)):
                            row_cells[1].text = f"{value:.3f}"
                        else:
                            row_cells[1].text = str(value)
            
            # 2.3 目标力值分析
            doc.add_heading('2.3 目标力值分析', level=2)
            target_analysis = analysis_data.get('target_analysis', [])
            if target_analysis and isinstance(target_analysis, list):
                target_table = doc.add_table(rows=1, cols=8)
                target_table.style = 'Table Grid'
                
                headers = ['目标力值(N)', '数据点数', '成功率(%)', '平均力值(N)', 
                          '平均偏差(N)', '标准差(N)', '最大偏差(N)', '绝对容差限制(N)']
                hdr_cells = target_table.rows[0].cells
                for i, header in enumerate(headers):
                    hdr_cells[i].text = header
                    hdr_cells[i].paragraphs[0].runs[0].font.bold = True
                
                for target in target_analysis:
                    if isinstance(target, dict):
                        row_cells = target_table.add_row().cells
                        row_cells[0].text = str(target.get('target_force', ''))
                        row_cells[1].text = str(target.get('数据点数', ''))
                        row_cells[2].text = f"{target.get('成功率_综合', 0):.1f}"
                        row_cells[3].text = f"{target.get('平均力值', 0):.2f}"
                        row_cells[4].text = f"{target.get('平均偏差_绝对', 0):.2f}"
                        row_cells[5].text = f"{target.get('标准差', 0):.2f}"
                        row_cells[6].text = f"{target.get('最大偏差_绝对', 0):.2f}"
                        row_cells[7].text = f"{target.get('绝对容差限制', 0):.2f}"
            
            # 2.4 过程能力分析
            doc.add_heading('2.4 过程能力分析', level=2)
            process_capability = analysis_data.get('process_capability', [])
            if process_capability and isinstance(process_capability, list):
                capability_table = doc.add_table(rows=1, cols=4)
                capability_table.style = 'Table Grid'
                
                headers = ['目标力值(N)', 'Cp', 'Cpk', '能力等级']
                hdr_cells = capability_table.rows[0].cells
                for i, header in enumerate(headers):
                    hdr_cells[i].text = header
                    hdr_cells[i].paragraphs[0].runs[0].font.bold = True
                
                for cap in process_capability:
                    if isinstance(cap, dict):
                        row_cells = capability_table.add_row().cells
                        row_cells[0].text = str(cap.get('target_force', ''))
                        row_cells[1].text = f"{cap.get('Cp', 0):.3f}"
                        row_cells[2].text = f"{cap.get('Cpk', 0):.3f}"
                        row_cells[3].text = cap.get('能力等级', '')
            
            # 2.5 位置区域分析（机器人压力测试特有）
            doc.add_heading('2.5 位置区域分析', level=2)
            multi_source = analysis_data.get('multi_source_variation_analysis', {})
            position_analysis = multi_source.get('performance_by_position', [])
            if position_analysis:
                position_table = doc.add_table(rows=1, cols=6)
                position_table.style = 'Table Grid'
                
                headers = ['位置区域', '目标力值(N)', '数据点数', '成功率(%)', '平均偏差(N)', '标准差(N)']
                hdr_cells = position_table.rows[0].cells
                for i, header in enumerate(headers):
                    hdr_cells[i].text = header
                    hdr_cells[i].paragraphs[0].runs[0].font.bold = True
                
                for pos in position_analysis:
                    if isinstance(pos, dict):
                        row_cells = position_table.add_row().cells
                        row_cells[0].text = pos.get('position_group', '')
                        row_cells[1].text = str(pos.get('target_force', ''))
                        row_cells[2].text = str(pos.get('数据点数', ''))
                        row_cells[3].text = f"{pos.get('成功率_综合', 0):.1f}"
                        row_cells[4].text = f"{pos.get('平均偏差_绝对', 0):.2f}"
                        row_cells[5].text = f"{pos.get('标准差', 0):.2f}"
            
            # 2.6 机器人一致性分析
            doc.add_heading('2.6 机器人一致性分析', level=2)
            robot_consistency = multi_source.get('robot_consistency_analysis', {})
            if robot_consistency:
                doc.add_paragraph("**力值重复性分析：**")
                force_repeatability = robot_consistency.get('force_repeatability', {})
                
                # 处理 force_repeatability 可能是列表的情况
                if force_repeatability:
                    if isinstance(force_repeatability, dict):
                        for target_force, cv in force_repeatability.items():
                            doc.add_paragraph(f"• 目标{target_force}N的变异系数: {cv:.2f}%")
                    elif isinstance(force_repeatability, list) and len(force_repeatability) > 0:
                        # 如果是列表，尝试处理每个元素
                        for item in force_repeatability:
                            if isinstance(item, dict):
                                target_force = item.get('target_force', '未知')
                                cv = item.get('cv', item.get('变异系数', 0))
                                doc.add_paragraph(f"• 目标{target_force}N的变异系数: {cv:.2f}%")
                    else:
                        doc.add_paragraph("• 力值重复性数据格式异常")
                
                doc.add_paragraph("\n**位置精度分析：**")
                position_accuracy = robot_consistency.get('position_accuracy', {})
                
                # 处理 position_accuracy 可能是列表的情况
                if position_accuracy:
                    if isinstance(position_accuracy, dict):
                        for coord, std in position_accuracy.items():
                            doc.add_paragraph(f"• {coord}坐标标准差: {std:.3f}mm")
                    elif isinstance(position_accuracy, list) and len(position_accuracy) > 0:
                        # 如果是列表，尝试处理每个元素
                        for item in position_accuracy:
                            if isinstance(item, dict):
                                coord = item.get('coordinate', item.get('坐标', '未知'))
                                std = item.get('std', item.get('标准差', 0))
                                doc.add_paragraph(f"• {coord}坐标标准差: {std:.3f}mm")
                    else:
                        doc.add_paragraph("• 位置精度数据格式异常")
            
            # 2.7 关键图表展示
            doc.add_heading('2.7 完整图表展示', level=2)
            
            # 检查图表目录是否存在
            charts_dir = Path(settings.CHARTS_DIR) / task_id
            logger.info(f"检查图表目录: {charts_dir}")
            
            if not charts_dir.exists():
                logger.warning(f"图表目录不存在: {charts_dir}")
                doc.add_paragraph("⚠️ 图表目录不存在，跳过图表展示")
                doc.add_paragraph(f"期望的图表目录路径: {charts_dir}")
                charts = []
            else:
                logger.info(f"图表目录存在，开始收集图表")
                charts = self._collect_charts(task_id, charts_dir)
                logger.info(f"收集到 {len(charts)} 个图表")
            
            if charts:
                # 按类别组织图表
                categories = {}
                for chart in charts:
                    category = chart['category']
                    if category not in categories:
                        categories[category] = []
                    categories[category].append(chart)
                
                # 显示所有图表，按类别分组
                category_order = ["基础分析", "偏差分析", "统计过程控制", "空间分析", "高级分析", "过程能力", "质量仪表盘", "误差分布分析", "多源变异分析"]
                
                chart_count = 0
                for category in category_order:
                    if category in categories and categories[category]:
                        doc.add_heading(f'{category}', level=3)
                        doc.add_paragraph(f"本类别包含 {len(categories[category])} 个图表：")
                        
                        # 显示该类别下的所有图表
                        for chart in categories[category]:
                            try:
                                chart_count += 1
                                
                                # 图表编号和标题
                                chart_title = doc.add_paragraph()
                                chart_title_run = chart_title.add_run(f"图表 {chart_count}: {chart['title']}")
                                chart_title_run.font.size = Pt(12)
                                chart_title_run.font.bold = True
                                
                                # 插入图表
                                if os.path.exists(chart['file_path']):
                                    # 添加图表
                                    doc.add_picture(chart['file_path'], width=Inches(6.0))
                                    
                                    # 图表描述
                                    desc_para = doc.add_paragraph()
                                    desc_run = desc_para.add_run(f"描述：{chart['description']}")
                                    desc_run.font.size = Pt(10)
                                    desc_run.font.italic = True
                                    
                                    # 图表解读说明
                                    if chart.get('interpretation'):
                                        interp_para = doc.add_paragraph()
                                        interp_run = interp_para.add_run("解读说明：")
                                        interp_run.font.size = Pt(10)
                                        interp_run.font.bold = True
                                        
                                        # 添加解读内容
                                        interpretation_text = chart['interpretation'].strip()
                                        interp_content = doc.add_paragraph()
                                        interp_content_run = interp_content.add_run(interpretation_text)
                                        interp_content_run.font.size = Pt(9)
                                        interp_content_run.font.color.rgb = RGBColor(64, 64, 64)  # 深灰色
                                        
                                else:
                                    doc.add_paragraph(f"⚠️ 图表文件未找到: {chart['file_path']}")
                                
                                # 添加分隔符
                                doc.add_paragraph('')
                                
                            except Exception as e:
                                logger.error(f"插入图表失败: {chart['title']} - {e}")
                                doc.add_paragraph(f"❌ 图表插入失败: {chart['title']}")
                
                # 添加图表汇总信息
                doc.add_paragraph('')
                summary_para = doc.add_paragraph()
                summary_run = summary_para.add_run(f"✅ 总计插入 {chart_count} 个分析图表")
                summary_run.font.size = Pt(12)
                summary_run.font.bold = True
                summary_run.font.color.rgb = RGBColor(0, 128, 0)  # 绿色
                
                # 添加图表类别统计
                doc.add_paragraph('')
                doc.add_paragraph("📊 图表类别统计：")
                for category in category_order:
                    if category in categories:
                        count = len(categories[category])
                        doc.add_paragraph(f"  • {category}: {count} 个图表")
            
            else:
                doc.add_paragraph("⚠️ 未找到任何图表文件")
            
            # ============================================================================
            # 第三部分：综合结论与建议
            # ============================================================================
            doc.add_page_break()
            doc.add_heading('第三部分：综合结论与建议', level=1)
            
            doc.add_heading('3.1 分析方法对比', level=2)
            doc.add_paragraph("""
本报告采用了两种互补的分析方法：

1. **DeepSeek AI智能分析**：利用人工智能的模式识别能力，对数据进行全面的质量评估和异常检测，提供专业的工业质量控制建议。

2. **R统计分析**：基于统计学原理，进行精确的数值计算、图表生成和数据建模，提供量化的分析结果。

两种方法的结合确保了分析结果的准确性、全面性和实用性。
            """)
            
            doc.add_heading('3.2 综合改进建议', level=2)
            recommendations = self._generate_recommendations(analysis_data)
            for i, rec in enumerate(recommendations, 1):
                doc.add_paragraph(f"{i}. {rec}")
            
            # 保存报告到两个位置
            # 1. temp/reports 目录 (临时下载)
            temp_reports_dir = Path(__file__).parent.parent.parent / "temp" / "reports"
            temp_reports_dir.mkdir(parents=True, exist_ok=True)
            temp_report_path = temp_reports_dir / f"comprehensive_analysis_report_{task_id}.docx"
            
            # 2. output/reports 目录 (永久存储)
            output_reports_dir = Path(settings.CHARTS_DIR).parent / "reports"
            output_reports_dir.mkdir(parents=True, exist_ok=True)
            output_report_path = output_reports_dir / f"comprehensive_analysis_report_{task_id}.docx"
            
            # 保存到临时目录
            doc.save(str(temp_report_path))
            logger.info(f"综合Word报告已保存到临时目录: {temp_report_path}")
            
            # 复制到输出目录
            import shutil
            shutil.copy2(temp_report_path, output_report_path)
            logger.info(f"综合Word报告已保存到输出目录: {output_report_path}")
            
            return str(temp_report_path)
            
        except Exception as e:
            logger.error(f"生成综合Word报告失败: {e}")
            import traceback
            logger.error(f"详细错误追踪: {traceback.format_exc()}")
            return None 

    def save_to_history(self, task_id: str, analysis_results: Dict[str, Any], original_filename: str):
        """将分析结果的关键信息保存到历史记录中"""
        try:
            history_path = Path(settings.HISTORY_DIR)
            history_path.mkdir(exist_ok=True)
            
            # 从分析结果中提取摘要信息
            # 使用 .get() 方法以避免因缺少键而引发错误
            overall_summary = analysis_results.get('overall_summary', {})
            if isinstance(overall_summary, list) and len(overall_summary) > 0:
                 # R toJSON 可能会将单行data.frame转为list of lists
                overall_summary = overall_summary[0]

            success_rate = overall_summary.get('综合成功率', 0)
            
            # 生成默认任务名：文件名 + 时间
            current_time = datetime.now(pytz.timezone('Asia/Shanghai'))
            time_str = current_time.strftime("%Y%m%d_%H%M%S")
            default_name = f"{Path(original_filename).stem}_{time_str}"
            
            history_record = {
                "id": task_id,
                "name": default_name,
                "original_filename": original_filename,
                "date": current_time.isoformat(),
                "created_at": current_time.isoformat(),
                "successRate": success_rate,
                "file_id": task_id  # 确保历史记录中有对文件的引用
            }
            
            history_file = history_path / f"{task_id}.json"
            with open(history_file, 'w', encoding='utf-8') as f:
                json.dump(history_record, f, ensure_ascii=False, indent=4)
            
            logger.info(f"任务 {task_id} 已成功保存到历史记录。")

        except Exception as e:
            logger.error(f"保存历史记录失败 (任务ID: {task_id}): {e}", exc_info=True)